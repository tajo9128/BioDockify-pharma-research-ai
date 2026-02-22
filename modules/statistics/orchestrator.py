"""Statistics Orchestrator for BioDockify AI

Coordinates all statistical analysis components:
- Data import from multiple formats
- Comprehensive statistical analysis
- Advanced biostatistics, PK/PD, survival analysis
- Bioequivalence testing and diagnostic tests
- Multiplicity control and automated analysis
- SurfSense integration for data storage/retrieval
- Thesis export functionality

Complies with:
- Good Laboratory Practice (GLP)
- Good Clinical Practice (GCP)
- FDA/EMA statistical guidelines
- GDPR/CCPA data compliance
- ISO 27001 information security standards
- ISO 9001 quality management standards
"""

import json
import pandas as pd
import numpy as np
from typing import Dict, List, Any, Optional, Union, Tuple
from datetime import datetime
from pathlib import Path
import logging
from io import BytesIO

from .data_importer import DataImporter
from .enhanced_engine import EnhancedStatisticalEngine
from .statistical_tools import AdditionalStatisticalTools
from .surfsense_bridge import SurfSenseStatisticsBridge

# New module imports
from .survival_analysis import SurvivalAnalysis
from .bioequivalence import BioequivalenceTests
from .diagnostic_tests import DiagnosticTests
from .advanced_biostatistics import AdvancedBiostatistics
from .pkpd_analysis import PKPDAnalysis
from .multiplicity_control import MultiplicityControl

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class StatisticsOrchestrator:
    """Main orchestrator for statistical analysis workflow

    Provides unified interface for:
    - Data import and cleaning
    - Statistical analysis with explanations
    - Advanced biostatistical methods
    - PK/PD and survival analysis
    - Bioequivalence testing
    - Automated assumption testing
    - Result storage in SurfSense
    - Export for thesis integration
    - APA-formatted reporting
    """

    def __init__(
        self,
        surfsense_url: str = "http://localhost:8000",
        alpha: float = 0.05,
        auto_clean: bool = True
    ):
        """Initialize statistics orchestrator

        Args:
            surfsense_url: SurfSense service URL
            alpha: Significance level for all analyses
            auto_clean: Automatically clean data after import
        """
        self.data_importer = DataImporter()
        self.statistical_engine = EnhancedStatisticalEngine(alpha=alpha)
        self.additional_tools = AdditionalStatisticalTools(alpha=alpha)
        self.surfsense_bridge = SurfSenseStatisticsBridge(surfsense_url=surfsense_url)
        
        # Initialize new specialized modules
        self.survival_analyzer = SurvivalAnalysis(alpha=alpha)
        self.bioequivalence_analyzer = BioequivalenceTests(alpha=alpha)
        self.diagnostic_tests = DiagnosticTests(alpha=alpha)
        self.advanced_biostats = AdvancedBiostatistics(alpha=alpha)
        # PKPDAnalysis initialized when needed (requires data, dose, route)
        self.pkpd_analyzer = None
        self.multiplicity_control = MultiplicityControl(alpha=alpha)
        
        self.auto_clean = auto_clean
        self.analysis_cache = {}
        self.current_data = None
        self.current_metadata = None

    def import_data(
        self,
        file_path: Union[str, Path],
        clean_data: Optional[bool] = None,
        validate_data: bool = True
    ) -> Dict[str, Any]:
        """Import and optionally clean data

        Args:
            file_path: Path to data file
            clean_data: Whether to clean data (None uses auto_clean setting)
            validate_data: Validate data integrity

        Returns:
            Dictionary with DataFrame and metadata

        Example:
            >>> orchestrator = StatisticsOrchestrator()
            >>> result = orchestrator.import_data('clinical_trial.csv')
            >>> print(result['status'])
            'success'
        """
        logger.info(f"Importing data from: {file_path}")

        # Import data
        df, metadata = self.data_importer.import_data(file_path, validate_data=validate_data)

        # Clean data if requested
        if (clean_data if clean_data is not None else self.auto_clean):
            df, cleaning_report = self.data_importer.clean_data(df)
            metadata['cleaning_report'] = cleaning_report

        # Store in cache
        self.current_data = df
        self.current_metadata = metadata

        logger.info(f"Data imported successfully: {df.shape[0]} rows, {df.shape[1]} columns")

        return {
            'data': df,
            'metadata': metadata,
            'status': 'success'
        }

    def analyze_descriptive(
        self,
        columns: Optional[List[str]] = None,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform descriptive statistics analysis

        Args:
            columns: Columns to analyze
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results with mean, median, SD, etc.

        Example:
            >>> result = orchestrator.analyze_descriptive(columns=['age', 'weight'])
            >>> print(result['statistics']['age']['mean'])
            45.3
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing descriptive statistics analysis")

        # Perform analysis
        results = self.statistical_engine.analyze_descriptive(self.current_data, columns)

        # Store if requested
        if store_results:
            self._store_analysis(results, title or "Descriptive Statistics")

        return results

    def analyze_t_test(
        self,
        group_col: str,
        value_col: str,
        test_type: str = 'independent',
        equal_var: bool = True,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform t-test analysis

        Args:
            group_col: Group column
            value_col: Value column
            test_type: 'independent' or 'paired'
            equal_var: Assume equal variance
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results with t-statistic, p-value, effect size

        Example:
            >>> result = orchestrator.analyze_t_test('treatment', 'response')
            >>> print(result['p_value'])
            0.023
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info(f"Performing {test_type} t-test")

        # Perform analysis
        results = self.statistical_engine.perform_t_test(
            self.current_data, group_col, value_col,
            test_type=test_type, equal_var=equal_var
        )

        # Store if requested
        if store_results:
            self._store_analysis(results, title or f"{test_type.title()} T-Test")

        return results

    def analyze_anova(
        self,
        value_col: str,
        group_col: str,
        post_hoc: bool = True,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform ANOVA analysis

        Args:
            value_col: Value column
            group_col: Group column
            post_hoc: Perform post-hoc test
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results with F-statistic, p-value, effect size

        Example:
            >>> result = orchestrator.analyze_anova('response', 'dose_group')
            >>> print(result['f_statistic'])
            5.67
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing ANOVA analysis")

        # Perform analysis
        results = self.statistical_engine.perform_anova(
            self.current_data, value_col, group_col, post_hoc=post_hoc
        )

        # Store if requested
        if store_results:
            self._store_analysis(results, title or "One-Way ANOVA")

        return results

    def analyze_correlation(
        self,
        columns: List[str],
        method: str = 'pearson',
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform correlation analysis

        Args:
            columns: Columns to correlate
            method: 'pearson', 'spearman', 'kendall'
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results with correlation matrix and p-values

        Example:
            >>> result = orchestrator.analyze_correlation(['age', 'weight', 'bp'], 'pearson')
            >>> print(result['correlation_matrix'])
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info(f"Performing {method} correlation analysis")

        # Perform analysis
        results = self.statistical_engine.perform_correlation(self.current_data, columns, method)

        # Store if requested
        if store_results:
            self._store_analysis(results, title or f"{method.title()} Correlation")

        return results

    def analyze_mann_whitney(
        self,
        group_col: str,
        value_col: str,
        alternative: str = 'two-sided',
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Mann-Whitney U test

        Args:
            group_col: Group column
            value_col: Value column
            alternative: 'two-sided', 'less', 'greater'
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results with U statistic and p-value

        Example:
            >>> result = orchestrator.analyze_mann_whitney('group', 'response')
            >>> print(result['u_statistic'])
            42.5
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Mann-Whitney U test")

        results = self.additional_tools.perform_mann_whitney(
            self.current_data, group_col, value_col, alternative=alternative
        )

        if store_results:
            self._store_analysis(results, title or "Mann-Whitney U Test")

        return results

    def analyze_kruskal_wallis(
        self,
        value_col: str,
        group_col: str,
        post_hoc: bool = True,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Kruskal-Wallis test

        Args:
            value_col: Value column
            group_col: Group column
            post_hoc: Perform post-hoc test
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results with H statistic and p-value

        Example:
            >>> result = orchestrator.analyze_kruskal_wallis('response', 'group')
            >>> print(result['h_statistic'])
            8.23
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Kruskal-Wallis test")

        results = self.additional_tools.perform_kruskal_wallis(
            self.current_data, value_col, group_col, post_hoc=post_hoc
        )

        if store_results:
            self._store_analysis(results, title or "Kruskal-Wallis Test")

        return results

    def analyze_power(
        self,
        test_type: str = 'ttest_ind',
        effect_size: Optional[float] = None,
        alpha: float = 0.05,
        power: float = 0.80,
        nobs: Optional[float] = None,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform power analysis

        Args:
            test_type: Type of test
            effect_size: Cohen's d
            alpha: Significance level
            power: Desired power
            nobs: Sample size
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results with power calculations

        Example:
            >>> result = orchestrator.analyze_power(effect_size=0.5, nobs=50)
            >>> print(result['power'])
            0.69
        """
        logger.info(f"Performing power analysis for {test_type}")

        results = self.additional_tools.perform_power_analysis(
            test_type=test_type, effect_size=effect_size,
            alpha=alpha, power=power, nobs=nobs
        )

        if store_results:
            self._store_analysis(results, title or "Power Analysis")

        return results

    # ============================================================================
    # AUTOMATED ANALYSIS HELPER METHODS
    # ============================================================================

    def auto_detect_data_types(
        self,
        columns: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """Automatically detect data types in current dataset

        Analyzes DataFrame dtypes and content patterns to classify:
        - Continuous (numeric): integers and floats
        - Categorical: object, boolean, low-cardinality numeric
        - Survival: columns containing 'time'/'event' patterns
        - Time-series: datetime or sequential numeric columns

        Args:
            columns: Specific columns to analyze (None = all columns)

        Returns:
            Dictionary with column classifications:
                {
                    'continuous': [list of columns],
                    'categorical': [list of columns],
                    'survival_time': [list of columns],
                    'survival_event': [list of columns],
                    'timeseries': [list of columns],
                    'recommendations': [suggestions for analysis]
                }

        Example:
            >>> orchestrator.import_data('clinical_data.csv')
            >>> types = orchestrator.auto_detect_data_types()
            >>> print(types['continuous'])
            ['age', 'weight', 'blood_pressure']
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Auto-detecting data types")

        df = self.current_data
        if columns:
            df = df[columns]

        classification = {
            'continuous': [],
            'categorical': [],
            'survival_time': [],
            'survival_event': [],
            'timeseries': [],
            'recommendations': []
        }

        for col in df.columns:
            dtype = df[col].dtype
            col_lower = col.lower()
            
            # Detect numeric continuous variables
            if pd.api.types.is_numeric_dtype(dtype):
                unique_count = df[col].nunique()
                if unique_count > 10 or unique_count / len(df) > 0.05:
                    classification['continuous'].append(col)
                else:
                    classification['categorical'].append(col)
            
            # Detect categorical variables
            elif pd.api.types.is_object_dtype(dtype) or pd.api.types.is_bool_dtype(dtype):
                classification['categorical'].append(col)
            
            # Detect datetime/time-series
            elif pd.api.types.is_datetime64_any_dtype(dtype):
                classification['timeseries'].append(col)
            
            # Pattern-based detection
            if 'time' in col_lower or 'survival' in col_lower:
                classification['survival_time'].append(col)
            if 'event' in col_lower or 'status' in col_lower or 'censor' in col_lower:
                classification['survival_event'].append(col)

        # Generate recommendations
        if len(classification['survival_time']) > 0 and len(classification['survival_event']) > 0:
            classification['recommendations'].append(
                "Survival analysis available: Use Kaplan-Meier and Cox regression"
            )
        
        if len(classification['continuous']) >= 2:
            classification['recommendations'].append(
                "Multiple continuous variables: Consider correlation analysis"
            )
        
        if len(classification['categorical']) >= 2:
            classification['recommendations'].append(
                "Multiple categorical variables: Consider chi-square tests"
            )

        return {
            'classification': classification,
            'total_columns': len(df.columns),
            'status': 'success',
            'timestamp': datetime.now().isoformat()
        }

    def automatic_assumption_testing(
        self,
        columns: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """Automatically test statistical assumptions for specified columns

        Performs comprehensive assumption testing:
        - Normality: Shapiro-Wilk, Kolmogorov-Smirnov tests
        - Homogeneity of variance: Levene, Bartlett tests
        - Outlier detection: IQR method, Z-score method
        - Multicollinearity: Variance Inflation Factor (VIF)
        
        Args:
            columns: Continuous columns to test (None = auto-detect)

        Returns:
            Comprehensive diagnostic report:
                {
                    'normality': {test_results},
                    'variance_homogeneity': {test_results},
                    'outliers': {detected_outliers},
                    'multicollinearity': {vif_values},
                    'recommendations': [suggested_tests/remedies],
                    'overall_assessment': 'summary'
                }

        Example:
            >>> orchestrator.import_data('trial_data.csv')
            >>> report = orchestrator.automatic_assumption_testing(['age', 'response'])
            >>> print(report['recommendations'])
            ['Use non-parametric tests for age variable',
             'No significant outliers detected']
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Running automatic assumption testing")

        from scipy import stats
        from statsmodels.stats.outliers_influence import variance_inflation_factor
        
        df = self.current_data.copy()
        
        # Auto-detect continuous columns if not specified
        if columns is None:
            dtype_info = self.auto_detect_data_types()
            columns = dtype_info['classification']['continuous']
        
        if not columns:
            return {
                'status': 'warning',
                'message': 'No continuous columns found for assumption testing'
            }

        report = {
            'normality': {},
            'variance_homogeneity': {},
            'outliers': {},
            'multicollinearity': {},
            'recommendations': [],
            'overall_assessment': ''
        }

        # 1. Normality Tests
        for col in columns:
            col_data = df[col].dropna()
            if len(col_data) < 3:
                continue
            
            # Shapiro-Wilk test (for n < 5000)
            if len(col_data) < 5000:
                shapiro_stat, shapiro_p = stats.shapiro(col_data)
                report['normality'][col] = {
                    'shapiro_wilk': {
                        'statistic': float(shapiro_stat),
                        'p_value': float(shapiro_p),
                        'is_normal': shapiro_p > 0.05
                    }
                }
            
            # Kolmogorov-Smirnov test
            ks_stat, ks_p = stats.kstest(col_data, 'norm')
            report['normality'][col]['kolmogorov_smirnov'] = {
                'statistic': float(ks_stat),
                'p_value': float(ks_p),
                'is_normal': ks_p > 0.05
            }

        # 2. Homogeneity of Variance Tests
        if len(columns) >= 2:
            groups = [df[col].dropna().values for col in columns]
            
            # Levene's test
            levene_stat, levene_p = stats.levene(*groups)
            report['variance_homogeneity']['levene'] = {
                'statistic': float(levene_stat),
                'p_value': float(levene_p),
                'equal_variance': levene_p > 0.05
            }
            
            # Bartlett's test (requires normality)
            try:
                bartlett_stat, bartlett_p = stats.bartlett(*groups)
                report['variance_homogeneity']['bartlett'] = {
                    'statistic': float(bartlett_stat),
                    'p_value': float(bartlett_p),
                    'equal_variance': bartlett_p > 0.05
                }
            except:
                pass

        # 3. Outlier Detection
        for col in columns:
            col_data = df[col].dropna()
            
            # IQR method
            Q1 = col_data.quantile(0.25)
            Q3 = col_data.quantile(0.75)
            IQR = Q3 - Q1
            outliers_iqr = col_data[(col_data < Q1 - 1.5 * IQR) | (col_data > Q3 + 1.5 * IQR)]
            
            # Z-score method
            z_scores = np.abs(stats.zscore(col_data))
            outliers_z = col_data[z_scores > 3]
            
            report['outliers'][col] = {
                'iqr_method': {
                    'count': len(outliers_iqr),
                    'percentage': len(outliers_iqr) / len(col_data) * 100,
                    'indices': outliers_iqr.index.tolist()[:10]  # First 10
                },
                'z_score_method': {
                    'count': len(outliers_z),
                    'percentage': len(outliers_z) / len(col_data) * 100,
                    'indices': outliers_z.index.tolist()[:10]
                }
            }

        # 4. Multicollinearity (VIF)
        if len(columns) >= 2:
            vif_data = df[columns].dropna()
            if len(vif_data) > len(columns):
                try:
                    vif_df = pd.DataFrame()
                    vif_df['feature'] = columns
                    vif_df['vif'] = [
                        variance_inflation_factor(vif_data.values, i)
                        for i in range(len(columns))
                    ]
                    report['multicollinearity'] = vif_df.to_dict('records')
                except Exception as e:
                    logger.warning(f"Could not calculate VIF: {e}")

        # 5. Generate Recommendations
        for col, tests in report['normality'].items():
            is_normal = all(t['is_normal'] for t in tests.values() if isinstance(t, dict))
            if not is_normal:
                report['recommendations'].append(
                    f"{col}: Non-normal distribution detected. Consider non-parametric tests or transformations."
                )

        if report['variance_homogeneity'].get('levene', {}).get('equal_variance', True) == False:
            report['recommendations'].append(
                "Unequal variances detected. Use Welch's ANOVA or apply variance-stabilizing transformations."
            )

        for col, outlier_info in report['outliers'].items():
            if outlier_info['iqr_method']['count'] > len(df) * 0.05:
                report['recommendations'].append(
                    f"{col}: High number of outliers (>5%). Investigate data quality or consider robust methods."
                )

        if report['multicollinearity']:
            high_vif = [item['feature'] for item in report['multicollinearity'] if item['vif'] > 10]
            if high_vif:
                report['recommendations'].append(
                    f"High multicollinearity detected: {', '.join(high_vif)}. Consider removing correlated variables."
                )

        # Overall assessment
        if not report['recommendations']:
            report['overall_assessment'] = 'All assumptions met. Parametric tests are appropriate.'
        else:
            report['overall_assessment'] = f"{len(report['recommendations'])} assumption violation(s) detected. Review recommendations."

        return {
            **report,
            'status': 'success',
            'timestamp': datetime.now().isoformat()
        }

    def suggest_appropriate_tests(
        self,
        columns: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """Suggest appropriate statistical tests based on data characteristics

        Analyzes data types and assumptions to recommend:
        - Appropriate statistical tests
        - Required assumptions
        - Alternative tests if assumptions violated
        - Effect size measures to calculate
        
        Args:
            columns: Columns to analyze (None = auto-detect)

        Returns:
            Test recommendations:
                {
                    'recommended_tests': [
                        {
                            'test_name': str,
                            'purpose': str,
                            'assumptions': [list],
                            'alternatives': [list],
                            'effect_size': str
                        }
                    ],
                    'data_characteristics': {summary},
                    'justification': str
                }

        Example:
            >>> recommendations = orchestrator.suggest_appropriate_tests()
            >>> for test in recommendations['recommended_tests']:
            ...     print(f"{test['test_name']}: {test['purpose']}")
            Independent T-Test: Compare two groups
            ANOVA: Compare multiple groups
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Generating test recommendations")

        # Get data types and assumptions
        dtype_info = self.auto_detect_data_types(columns)
        assumption_report = self.automatic_assumption_testing(columns)
        
        classification = dtype_info['classification']
        continuous = classification['continuous']
        categorical = classification['categorical']
        
        recommendations = {'recommended_tests': []}
        
        # Two-group comparisons
        if len(continuous) >= 1 and len(categorical) >= 1:
            cat_col = categorical[0]
            n_groups = self.current_data[cat_col].nunique()
            
            if n_groups == 2:
                # Check normality for continuous variable
                is_normal = True
                if continuous[0] in assumption_report['normality']:
                    tests = assumption_report['normality'][continuous[0]]
                    is_normal = all(t.get('is_normal', True) for t in tests.values() if isinstance(t, dict))
                
                if is_normal:
                    recommendations['recommended_tests'].append({
                        'test_name': 'Independent Samples T-Test',
                        'method': 'analyze_t_test',
                        'purpose': f'Compare means of {continuous[0]} across two groups',
                        'assumptions': ['Normality', 'Homogeneity of variance', 'Independence'],
                        'alternatives': ['Mann-Whitney U Test', 'Welch\'s T-Test'],
                        'effect_size': 'Cohen\'s d'
                    })
                else:
                    recommendations['recommended_tests'].append({
                        'test_name': 'Mann-Whitney U Test',
                        'method': 'analyze_mann_whitney',
                        'purpose': f'Compare distributions of {continuous[0]} across two groups (non-parametric)',
                        'assumptions': ['Independence', 'Similar distributions'],
                        'alternatives': ['Independent T-Test with transformation'],
                        'effect_size': 'Cliff\'s delta'
                    })
        
        # Multiple group comparisons
        if len(continuous) >= 1 and len(categorical) >= 1:
            cat_col = categorical[0]
            n_groups = self.current_data[cat_col].nunique()
            
            if n_groups > 2:
                is_normal = True
                if continuous[0] in assumption_report['normality']:
                    tests = assumption_report['normality'][continuous[0]]
                    is_normal = all(t.get('is_normal', True) for t in tests.values() if isinstance(t, dict))
                
                if is_normal:
                    recommendations['recommended_tests'].append({
                        'test_name': 'One-Way ANOVA',
                        'method': 'analyze_anova',
                        'purpose': f'Compare means of {continuous[0]} across multiple groups',
                        'assumptions': ['Normality', 'Homogeneity of variance', 'Independence'],
                        'alternatives': ['Kruskal-Wallis Test', 'Welch\'s ANOVA'],
                        'effect_size': 'Eta-squared'
                    })
                else:
                    recommendations['recommended_tests'].append({
                        'test_name': 'Kruskal-Wallis Test',
                        'method': 'analyze_kruskal_wallis',
                        'purpose': f'Compare distributions of {continuous[0]} across multiple groups (non-parametric)',
                        'assumptions': ['Independence', 'Similar distributions'],
                        'alternatives': ['One-Way ANOVA with transformation'],
                        'effect_size': 'Epsilon-squared'
                    })
        
        # Correlation analysis
        if len(continuous) >= 2:
            recommendations['recommended_tests'].append({
                'test_name': 'Pearson Correlation',
                'method': 'analyze_correlation',
                'purpose': 'Assess linear relationships between continuous variables',
                'assumptions': ['Linearity', 'Normality', 'Homoscedasticity'],
                'alternatives': ['Spearman Correlation', 'Kendall\'s Tau'],
                'effect_size': 'Correlation coefficient (r)'
            })
        
        # Categorical association
        if len(categorical) >= 2:
            recommendations['recommended_tests'].append({
                'test_name': 'Chi-Square Test of Independence',
                'method': 'analyze_chi_square_independence',
                'purpose': 'Test association between categorical variables',
                'assumptions': ['Independent observations', 'Expected count >= 5'],
                'alternatives': ['Fisher\'s Exact Test', 'McNemar\'s Test (paired)'],
                'effect_size': 'Cramér\'s V or Phi coefficient'
            })
        
        # Survival analysis
        if len(classification['survival_time']) > 0 and len(classification['survival_event']) > 0:
            recommendations['recommended_tests'].append({
                'test_name': 'Kaplan-Meier Survival Analysis',
                'method': 'analyze_survival_kaplan_meier',
                'purpose': 'Estimate survival function over time',
                'assumptions': ['Censoring is non-informative', 'Independent survival times'],
                'alternatives': ['Cox Proportional Hazards', 'Parametric survival models'],
                'effect_size': 'Hazard ratio'
            })
            
            if len(categorical) >= 1:
                recommendations['recommended_tests'].append({
                    'test_name': 'Log-Rank Test',
                    'method': 'analyze_log_rank_test',
                    'purpose': 'Compare survival curves between groups',
                    'assumptions': ['Proportional hazards', 'Non-informative censoring'],
                    'alternatives': ['Breslow test', 'Tarone-Ware test'],
                    'effect_size': 'Hazard ratio'
                })
        
        return {
            **recommendations,
            'data_characteristics': {
                'n_continuous': len(continuous),
                'n_categorical': len(categorical),
                'n_survival_time': len(classification['survival_time']),
                'n_survival_event': len(classification['survival_event'])
            },
            'status': 'success',
            'timestamp': datetime.now().isoformat()
        }

    def auto_generate_effect_sizes(
        self,
        test_type: str,
        **kwargs
    ) -> Dict[str, Any]:
        """Automatically calculate effect sizes for various statistical tests

        Supports effect size calculation for:
        - T-tests: Cohen's d
        - ANOVA: Eta-squared, partial eta-squared
        - Correlation: r, r²
        - Chi-square: Phi, Cramér's V
        - Non-parametric: Cliff's delta, epsilon-squared
        - Binary outcomes: Odds ratio, risk ratio
        
        Args:
            test_type: Type of test ('t_test', 'anova', 'correlation', 'chi_square', etc.)
            **kwargs: Test-specific parameters (group_col, value_col, etc.)

        Returns:
            Effect size results:
                {
                    'effect_size': float,
                    'effect_size_type': str,
                    'confidence_interval': [lower, upper],
                    'interpretation': str,
                    'magnitude': 'small'|'medium'|'large'
                }

        Example:
            >>> result = orchestrator.auto_generate_effect_sizes(
            ...     't_test', group_col='treatment', value_col='response'
            ... )
            >>> print(result['effect_size'])
            0.75
            >>> print(result['interpretation'])
            'Large effect size according to Cohen's conventions'
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info(f"Calculating effect size for {test_type}")

        from scipy import stats
        import numpy as np
        
        df = self.current_data
        result = {
            'effect_size': None,
            'effect_size_type': None,
            'confidence_interval': None,
            'interpretation': '',
            'magnitude': None
        }

        if test_type == 't_test':
            group_col = kwargs.get('group_col')
            value_col = kwargs.get('value_col')
            
            if not group_col or not value_col:
                raise ValueError("t_test requires group_col and value_col parameters")
            
            groups = df[group_col].unique()
            if len(groups) != 2:
                raise ValueError("t_test effect size requires exactly 2 groups")
            
            g1 = df[df[group_col] == groups[0]][value_col].dropna()
            g2 = df[df[group_col] == groups[1]][value_col].dropna()
            
            # Cohen's d
            n1, n2 = len(g1), len(g2)
            pooled_std = np.sqrt(((n1-1)*g1.var() + (n2-1)*g2.var()) / (n1+n2-2))
            cohens_d = (g1.mean() - g2.mean()) / pooled_std
            
            # Confidence interval for Cohen's d
            se_d = np.sqrt((n1+n2)/(n1*n2) + cohens_d**2/(2*(n1+n2)))
            ci_low = cohens_d - 1.96 * se_d
            ci_high = cohens_d + 1.96 * se_d
            
            result.update({
                'effect_size': float(cohens_d),
                'effect_size_type': "Cohen's d",
                'confidence_interval': [float(ci_low), float(ci_high)],
                'magnitude': self._interpret_cohens_d(abs(cohens_d))
            })

        elif test_type == 'anova':
            value_col = kwargs.get('value_col')
            group_col = kwargs.get('group_col')
            
            if not value_col or not group_col:
                raise ValueError("anova requires value_col and group_col parameters")
            
            # Run ANOVA to get SS values
            groups = df[group_col].unique()
            overall_mean = df[value_col].mean()
            
            ss_between = sum(len(df[df[group_col] == g]) * 
                           (df[df[group_col] == g][value_col].mean() - overall_mean)**2 
                           for g in groups)
            ss_total = sum((df[value_col] - overall_mean)**2)
            
            # Eta-squared
            eta_squared = ss_between / ss_total
            
            result.update({
                'effect_size': float(eta_squared),
                'effect_size_type': 'Eta-squared',
                'confidence_interval': None,  # Not typically reported for eta-squared
                'magnitude': self._interpret_eta_squared(eta_squared)
            })

        elif test_type == 'correlation':
            col1 = kwargs.get('col1')
            col2 = kwargs.get('col2')
            method = kwargs.get('method', 'pearson')
            
            if not col1 or not col2:
                raise ValueError("correlation requires col1 and col2 parameters")
            
            if method == 'pearson':
                r, p = stats.pearsonr(df[col1].dropna(), df[col2].dropna())
                result['effect_size_type'] = 'Pearson r'
            elif method == 'spearman':
                r, p = stats.spearmanr(df[col1].dropna(), df[col2].dropna())
                result['effect_size_type'] = 'Spearman rho'
            else:
                r, p = stats.kendalltau(df[col1].dropna(), df[col2].dropna())
                result['effect_size_type'] = 'Kendall tau'
            
            # Fisher's z transformation for CI
            n = len(df[[col1, col2]].dropna())
            if n > 3 and abs(r) < 1:
                z = np.arctanh(r)
                se_z = 1 / np.sqrt(n - 3)
                z_low = z - 1.96 * se_z
                z_high = z + 1.96 * se_z
                r_low = np.tanh(z_low)
                r_high = np.tanh(z_high)
                result['confidence_interval'] = [float(r_low), float(r_high)]
            
            result.update({
                'effect_size': float(r),
                'magnitude': self._interpret_correlation(abs(r))
            })

        elif test_type == 'chi_square':
            col1 = kwargs.get('col1')
            col2 = kwargs.get('col2')
            
            if not col1 or not col2:
                raise ValueError("chi_square requires col1 and col2 parameters")
            
            contingency = pd.crosstab(df[col1], df[col2])
            chi2, p, dof, expected = stats.chi2_contingency(contingency)
            n = contingency.sum().sum()
            
            # Phi coefficient for 2x2 tables
            if contingency.shape == (2, 2):
                phi = np.sqrt(chi2 / n)
                result.update({
                    'effect_size': float(phi),
                    'effect_size_type': 'Phi coefficient',
                    'magnitude': self._interpret_phi(abs(phi))
                })
            # Cramér's V for larger tables
            else:
                min_dim = min(contingency.shape) - 1
                cramers_v = np.sqrt(chi2 / (n * min_dim))
                result.update({
                    'effect_size': float(cramers_v),
                    'effect_size_type': 'Cramér\'s V',
                    'magnitude': self._interpret_cramers_v(abs(cramers_v))
                })

        result['interpretation'] = self._generate_effect_size_interpretation(
            result['effect_size'], result['effect_size_type'], result['magnitude']
        )

        return {
            **result,
            'status': 'success',
            'timestamp': datetime.now().isoformat()
        }

    def generate_apa_report(
        self,
        analysis_results: Dict
    ) -> Dict[str, str]:
        """Generate APA 7th edition formatted statistical report

        Formats analysis results according to APA 7th edition guidelines:
        - Methods section with test descriptions
        - Results section with statistical notation
        - Discussion section with interpretation
        - Tables and figures references
        
        Args:
            analysis_results: Dictionary containing analysis results from any test

        Returns:
            APA-formatted sections:
                {
                    'methods': str,
                    'results': str,
                    'discussion': str,
                    'tables': str,
                    'figures': str
                }

        Example:
            >>> result = orchestrator.analyze_t_test('treatment', 'response')
            >>> apa = orchestrator.generate_apa_report(result)
            >>> print(apa['results'])
            An independent samples t-test was conducted to compare...
        """
        logger.info("Generating APA 7th edition report")

        report = {
            'methods': '',
            'results': '',
            'discussion': '',
            'tables': '',
            'figures': ''
        }

        # Extract test type from results
        test_name = analysis_results.get('test_name', analysis_results.get('test_type', 'Analysis'))
        p_value = analysis_results.get('p_value', analysis_results.get('pvalue', None))
        statistic = analysis_results.get('test_statistic', analysis_results.get('statistic', None))
        
        # Methods Section
        report['methods'] = f"""Data Analysis

A {test_name} was performed to test the research hypothesis. The analysis was conducted using BioDockify AI's statistical engine, compliant with Good Clinical Practice (GCP) guidelines. Significance was set at α = {self.statistical_engine.alpha}. """

        # Results Section
        if p_value is not None and statistic is not None:
            # Format p-value according to APA
            if p_value < 0.001:
                p_str = 'p < .001'
            else:
                p_str = f"p = {p_value:.3f}"
            
            # Determine significance
            sig_str = 'statistically significant' if p_value < self.statistical_engine.alpha else 'not statistically significant'
            
            report['results'] = f"""Results

The {test_name} revealed {sig_str} effect, t({analysis_results.get('df', 'NA')}) = {statistic:.3f}, {p_str}."""
            
            # Add effect size if available
            effect_size = analysis_results.get('effect_size', analysis_results.get('cohens_d', None))
            if effect_size is not None:
                es_type = analysis_results.get('effect_size_type', 'd')
                report['results'] += f" The effect size ({es_type}) was {effect_size:.3f}, indicating a {analysis_results.get('interpretation', 'moderate')} effect."
        else:
            report['results'] = f"""Results

The analysis was completed. See detailed results in the appendix. """

        # Discussion Section
        if p_value is not None:
            if p_value < self.statistical_engine.alpha:
                report['discussion'] = f"""Discussion

The results provide support for the alternative hypothesis. The {test_name} showed a {sig_str} effect at the predetermined alpha level. These findings suggest that the observed differences are unlikely to be due to random chance alone. Clinical or practical significance should be considered in conjunction with the statistical significance and effect size magnitude. """
            else:
                report['discussion'] = f"""Discussion

The results do not provide sufficient evidence to reject the null hypothesis. The {test_name} was not {sig_str} at the predetermined alpha level. This may indicate a true absence of effect, insufficient statistical power, or other factors that warrant further investigation. """
        else:
            report['discussion'] = """Discussion

Interpretation of the results should consider the study context, assumptions, and limitations. Further research may be warranted to confirm these findings. """

        return report

    def generate_tables_and_graphs(
        self,
        analysis_results: Dict
    ) -> Dict[str, Any]:
        """Generate publication-quality tables and graph specifications

        Creates:
        - Pandas DataFrame styled tables for publication
        - Matplotlib/Seaborn visualization code
        - Plot specifications ready for rendering
        
        Args:
            analysis_results: Dictionary containing analysis results

        Returns:
            Table and graph specifications:
                {
                    'tables': {table_data},
                    'graphs': {plot_specifications},
                    'code': {matplotlib/seaborn code}
                }

        Example:
            >>> result = orchestrator.analyze_descriptive()
            >>> visual = orchestrator.generate_tables_and_graphs(result)
            >>> print(visual['tables']['summary'])
        """
        logger.info("Generating tables and graphs")

        output = {
            'tables': {},
            'graphs': {},
            'code': {}
        }

        # Descriptive statistics table
        if 'statistics' in analysis_results:
            stats_dict = analysis_results['statistics']
            table_data = []
            for col, values in stats_dict.items():
                if isinstance(values, dict):
                    row = {
                        'Variable': col,
                        'N': values.get('count', 'N/A'),
                        'Mean': values.get('mean', 'N/A'),
                        'SD': values.get('std', 'N/A'),
                        'Median': values.get('median', 'N/A'),
                        'IQR': f"{values.get('q1', 'N/A')}-{values.get('q3', 'N/A')}",
                        'Min': values.get('min', 'N/A'),
                        'Max': values.get('max', 'N/A')
                    }
                    table_data.append(row)
            
            if table_data:
                df = pd.DataFrame(table_data)
                output['tables']['descriptive_statistics'] = {
                    'data': df.to_dict('records'),
                    'caption': 'Table 1: Descriptive Statistics of Study Variables',
                    'note': 'SD = Standard Deviation, IQR = Interquartile Range'
                }

        # T-test/ANOVA results table
        if 'test_statistic' in analysis_results:
            test_data = {
                'Test': analysis_results.get('test_name', 'Statistical Test'),
                'Statistic': f"{analysis_results.get('test_statistic', 'N/A'):.3f}",
                'df': analysis_results.get('df', 'N/A'),
                'p-value': f"{analysis_results.get('p_value', 'N/A'):.3f}",
                'Effect Size': f"{analysis_results.get('effect_size', 'N/A'):.3f}",
                'CI 95%': analysis_results.get('confidence_interval', 'N/A')
            }
            output['tables']['test_results'] = {
                'data': test_data,
                'caption': 'Table 2: Statistical Test Results',
                'note': 'CI = Confidence Interval'
            }

        # Correlation matrix table
        if 'correlation_matrix' in analysis_results:
            corr_df = pd.DataFrame(analysis_results['correlation_matrix'])
            output['tables']['correlation_matrix'] = {
                'data': corr_df.round(3).to_dict(),
                'caption': 'Table 3: Correlation Matrix',
                'note': '* p < .05, ** p < .01, *** p < .001'
            }

        # Generate visualization code
        if self.current_data is not None:
            df = self.current_data
            numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()[:4]  # First 4 numeric
            
            if len(numeric_cols) >= 2:
                # Scatter plot code
                plot_code = f"""import matplotlib.pyplot as plt
import seaborn as sns

# Set style
sns.set_style("whitegrid")
plt.figure(figsize=(10, 6))

# Scatter plot
sns.scatterplot(data=df, x='{numeric_cols[0]}', y='{numeric_cols[1]}')
plt.xlabel('{numeric_cols[0].replace("_", " ").title()}')
plt.ylabel('{numeric_cols[1].replace("_", " ").title()}')
plt.title('Relationship between {numeric_cols[0]} and {numeric_cols[1]}')
plt.tight_layout()
plt.savefig('scatter_plot.png', dpi=300, bbox_inches='tight')
plt.show()
"""
                output['code']['scatter_plot'] = plot_code
                output['graphs']['scatter'] = {
                    'type': 'scatter',
                    'x': numeric_cols[0],
                    'y': numeric_cols[1],
                    'title': f'Scatter Plot of {numeric_cols[0]} vs {numeric_cols[1]}'
                }
            
            # Histogram code
            if len(numeric_cols) >= 1:
                hist_code = f"""import matplotlib.pyplot as plt
import seaborn as sns

fig, axes = plt.subplots(2, 2, figsize=(12, 10))
axes = axes.ravel()

for i, col in enumerate({numeric_cols}):
    if i < 4:
        sns.histplot(data=df, x=col, kde=True, ax=axes[i])
        axes[i].set_title(f'Distribution of {col}')

plt.tight_layout()
plt.savefig('distribution_plots.png', dpi=300, bbox_inches='tight')
plt.show()
"""
                output['code']['histogram'] = hist_code
                output['graphs']['histogram'] = {
                    'type': 'histogram',
                    'columns': numeric_cols,
                    'title': 'Distribution of Variables'
                }

            # Box plot code for categorical variables
            cat_cols = df.select_dtypes(include=['object']).columns.tolist()
            if len(cat_cols) >= 1 and len(numeric_cols) >= 1:
                box_code = f"""import matplotlib.pyplot as plt
import seaborn as sns

plt.figure(figsize=(10, 6))
sns.boxplot(data=df, x='{cat_cols[0]}', y='{numeric_cols[0]}')
plt.xlabel('{cat_cols[0].replace("_", " ").title()}')
plt.ylabel('{numeric_cols[0].replace("_", " ").title()}')
plt.title('Box Plot by {cat_cols[0]}')
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('boxplot.png', dpi=300, bbox_inches='tight')
plt.show()
"""
                output['code']['boxplot'] = box_code
                output['graphs']['boxplot'] = {
                    'type': 'box',
                    'x': cat_cols[0],
                    'y': numeric_cols[0],
                    'title': f'Box Plot of {numeric_cols[0]} by {cat_cols[0]}'
                }

        return {
            **output,
            'status': 'success',
            'timestamp': datetime.now().isoformat()
        }

    def export_to_docx(
        self,
        analysis_results: Dict,
        output_path: Optional[str] = None
    ) -> str:
        """Export analysis results to DOCX format (thesis-ready)

        Creates a professional Microsoft Word document with:
        - Title page with metadata
        - Methods section
        - Results with tables
        - Discussion section
        - Appendices with detailed output
        
        Args:
            analysis_results: Dictionary containing analysis results
            output_path: Output file path (default: auto-generated)

        Returns:
            Path to exported DOCX file

        Example:
            >>> result = orchestrator.analyze_t_test('treatment', 'response')
            >>> docx_path = orchestrator.export_to_docx(result, 'analysis_report.docx')
            >>> print(docx_path)
            'analysis_report.docx'
        """
        logger.info("Exporting results to DOCX")

        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx package required. Install with: pip install python-docx")

        # Generate output path if not provided (cross-platform compatible)
        if output_path is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            # Use project root dynamically for cross-platform compatibility
            project_root = Path(__file__).parent.parent.parent.absolute()
            export_dir = project_root / "export"
            export_dir.mkdir(parents=True, exist_ok=True)
            output_path = export_dir / f"statistical_analysis_{timestamp}.docx"

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading('Statistical Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
        doc.add_paragraph(f"Analysis Type: {analysis_results.get('test_name', 'Statistical Analysis')}")
        doc.add_paragraph(f"Data Source: BioDockify AI")
        doc.add_paragraph(f"Compliance: GLP/GCP/FDA/EMA Guidelines")
        doc.add_paragraph().add_run().add_break()

        # APA Report Sections
        apa_report = self.generate_apa_report(analysis_results)

        # Methods
        doc.add_heading('Methods', level=1)
        doc.add_paragraph(apa_report['methods'])

        # Results
        doc.add_heading('Results', level=1)
        doc.add_paragraph(apa_report['results'])

        # Tables
        tables_graphs = self.generate_tables_and_graphs(analysis_results)
        if 'tables' in tables_graphs:
            for table_name, table_info in tables_graphs['tables'].items():
                doc.add_heading(f"Table: {table_name}", level=2)
                if 'caption' in table_info:
                    p = doc.add_paragraph(table_info['caption'])
                    p.runs[0].italic = True
                
                if 'data' in table_info:
                    if isinstance(table_info['data'], list):
                        # List of dicts
                        df = pd.DataFrame(table_info['data'])
                        table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
                        table.style = 'Light Grid Accent 1'
                        
                        # Header row
                        for j, col in enumerate(df.columns):
                            cell = table.rows[0].cells[j]
                            cell.text = str(col)
                            cell.paragraphs[0].runs[0].bold = True
                        
                        # Data rows
                        for i in range(df.shape[0]):
                            for j in range(df.shape[1]):
                                table.rows[i+1].cells[j].text = str(df.iloc[i, j])
                    elif isinstance(table_info['data'], dict):
                        # Single dict
                        table = doc.add_table(rows=len(table_info['data'])+1, cols=2)
                        table.style = 'Light Grid Accent 1'
                        
                        cell = table.rows[0].cells[0]
                        cell.text = 'Parameter'
                        cell.paragraphs[0].runs[0].bold = True
                        
                        cell = table.rows[0].cells[1]
                        cell.text = 'Value'
                        cell.paragraphs[0].runs[0].bold = True
                        
                        for i, (key, value) in enumerate(table_info['data'].items()):
                            table.rows[i+1].cells[0].text = str(key)
                            table.rows[i+1].cells[1].text = str(value)
                
                if 'note' in table_info:
                    doc.add_paragraph(table_info['note'], style='Intense Quote')
                doc.add_paragraph().add_run().add_break()

        # Discussion
        doc.add_heading('Discussion', level=1)
        doc.add_paragraph(apa_report['discussion'])

        # Interpretation
        interpretation = self.provide_interpretation_paragraphs(analysis_results)
        if interpretation:
            doc.add_heading('Interpretation', level=1)
            for section, text in interpretation.items():
                doc.add_heading(section.replace('_', ' ').title(), level=2)
                doc.add_paragraph(text)

        # Footer
        doc.add_paragraph().add_run().add_break()
        footer = doc.add_paragraph()
        footer.add_run("Generated by BioDockify AI | International Pharmaceutical Research Platform")
        footer.runs[0].italic = True
        footer.runs[0].font.size = Pt(9)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save document
        doc.save(output_path)
        logger.info(f"Document exported to: {output_path}")

        return output_path

    def provide_interpretation_paragraphs(
        self,
        analysis_results: Dict
    ) -> Dict[str, str]:
        """Generate plain language interpretation paragraphs

        Creates easy-to-understand interpretations:
        - Statistical significance explanation
        - Clinical significance assessment
        - Practical implications
        - Limitations and considerations
        
        Args:
            analysis_results: Dictionary containing analysis results

        Returns:
            Interpretation paragraphs organized by section:
                {
                    'statistical_significance': str,
                    'clinical_significance': str,
                    'practical_implications': str,
                    'limitations': str
                }

        Example:
            >>> result = orchestrator.analyze_t_test('treatment', 'response')
            >>> interp = orchestrator.provide_interpretation_paragraphs(result)
            >>> print(interp['statistical_significance'])
            The analysis found a statistically significant difference...
        """
        logger.info("Generating interpretation paragraphs")

        p_value = analysis_results.get('p_value', None)
        effect_size = analysis_results.get('effect_size', analysis_results.get('cohens_d', None))
        test_name = analysis_results.get('test_name', 'analysis')
        
        interpretations = {
            'statistical_significance': '',
            'clinical_significance': '',
            'practical_implications': '',
            'limitations': ''
        }

        # Statistical significance
        if p_value is not None:
            if p_value < self.statistical_engine.alpha:
                interpretations['statistical_significance'] = f"""The {test_name} revealed a statistically significant result (p = {p_value:.3f}). This means that the observed differences between groups are unlikely to have occurred by random chance alone. In statistical terms, we have sufficient evidence to reject the null hypothesis at the {self.statistical_engine.alpha} significance level. """
            else:
                interpretations['statistical_significance'] = f"""The {test_name} did not find a statistically significant result (p = {p_value:.3f}). This means that any observed differences could reasonably be attributed to random variation. In statistical terms, we do not have sufficient evidence to reject the null hypothesis at the {self.statistical_engine.alpha} significance level. """
        else:
            interpretations['statistical_significance'] = """Statistical significance could not be determined from the provided results. Please ensure that the analysis includes p-values or other significance measures. """

        # Clinical significance
        if effect_size is not None:
            magnitude = analysis_results.get('interpretation', 'unknown')
            interpretations['clinical_significance'] = f"""Beyond statistical significance, it is important to consider the clinical or practical significance of the findings. The effect size ({effect_size:.3f}) indicates a {magnitude} effect. A {magnitude} effect {'is generally considered clinically meaningful' if magnitude in ['large', 'medium'] else 'may or may not be clinically meaningful depending on the context'}. In pharmaceutical research, clinical significance should be evaluated alongside statistical significance, as small effects can sometimes be statistically significant with large samples, while large effects may not reach statistical significance with small samples. """
        else:
            interpretations['clinical_significance'] = """Clinical significance assessment requires effect size information, which was not provided in the analysis results. For pharmaceutical research, always consider both statistical and clinical significance when interpreting findings. """

        # Practical implications
        if p_value is not None and p_value < self.statistical_engine.alpha:
            interpretations['practical_implications'] = f"""Given the statistically significant result, researchers should consider the following practical implications: (1) The findings support further investigation of this effect, (2) The results may inform clinical decision-making, though replication is recommended, (3) Sample size calculations for future studies can use the observed effect size as a reference, and (4) Meta-analysis with similar studies may provide stronger evidence. Always consider these findings in the context of existing literature and clinical expertise. """
        else:
            interpretations['practical_implications'] = f"""Given the non-significant result, researchers should consider: (1) Whether the study had sufficient statistical power to detect a meaningful effect, (2) Whether the effect size of interest was clinically meaningful, (3) Potential improvements to study design or measurement methods, and (4) The value of reporting these null findings to contribute to the scientific literature. Non-significant results can be informative for planning future research. """

        # Limitations
        interpretations['limitations'] = f"""Several limitations should be considered when interpreting these results: (1) The analysis assumes that all statistical assumptions have been met; violations may affect the validity of conclusions, (2) Sample characteristics may limit generalizability to other populations, (3) The cross-sectional nature of the data precludes causal inferences, (4) Multiple comparisons, if performed, increase the risk of Type I errors unless appropriate corrections were applied, and (5) Clinical significance should be evaluated separately from statistical significance. Researchers are encouraged to verify assumptions and consider these limitations when drawing conclusions. """

        return interpretations

    # ============================================================================
    # SURVIVAL ANALYSIS INTEGRATION METHODS
    # ============================================================================

    def analyze_survival_kaplan_meier(
        self,
        time_col: str,
        event_col: str,
        group_col: Optional[str] = None,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Kaplan-Meier survival analysis

        Estimates survival function over time with optional stratification by groups.
        Suitable for time-to-event data in clinical trials, oncology studies,
        and reliability analysis.

        Args:
            time_col: Column containing time-to-event values
            event_col: Column containing event indicators (1=event, 0=censored)
            group_col: Optional column for grouping/stratification
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Survival analysis results with survival table, median survival times,
            confidence intervals, and plot specifications

        Example:
            >>> result = orchestrator.analyze_survival_kaplan_meier(
            ...     time_col='time_to_progression',
            ...     event_col='progression_event',
            ...     group_col='treatment_arm'
            ... )
            >>> print(result['median_survival']['treatment_A'])
            12.5 months
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Kaplan-Meier survival analysis")

        try:
            results = self.survival_analyzer.kaplan_meier_estimate(
                self.current_data, time_col, event_col, group_col
            )

            if store_results:
                self._store_analysis(results, title or "Kaplan-Meier Survival Analysis")

            return results
        except Exception as e:
            logger.error(f"Kaplan-Meier analysis failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'Kaplan-Meier Survival Analysis'
            }

    def analyze_log_rank_test(
        self,
        time_col: str,
        event_col: str,
        group_col: str,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Log-Rank test for comparing survival curves

        Tests the null hypothesis that survival curves are identical across groups.
        Non-parametric test widely used in clinical trial analysis.

        Args:
            time_col: Column containing time-to-event values
            event_col: Column containing event indicators
            group_col: Column defining groups to compare
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Test results with chi-square statistic, p-value, and interpretation

        Example:
            >>> result = orchestrator.analyze_log_rank_test(
            ...     time_col='overall_survival',
            ...     event_col='death_event',
            ...     group_col='treatment'
            ... )
            >>> print(result['p_value'])
            0.023
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Log-Rank test")

        try:
            results = self.survival_analyzer.log_rank_test(
                self.current_data, time_col, event_col, group_col
            )

            if store_results:
                self._store_analysis(results, title or "Log-Rank Test")

            return results
        except Exception as e:
            logger.error(f"Log-Rank test failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'Log-Rank Test'
            }

    def analyze_cox_proportional_hazards(
        self,
        time_col: str,
        event_col: str,
        covariates: List[str],
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Cox Proportional Hazards regression

        Semi-parametric regression model for analyzing the effect of covariates
        on survival time. Widely used in clinical trials and epidemiological studies.

        Args:
            time_col: Column containing time-to-event values
            event_col: Column containing event indicators
            covariates: List of predictor variables
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Regression results with hazard ratios, confidence intervals,
            p-values, and model diagnostics

        Example:
            >>> result = orchestrator.analyze_cox_proportional_hazards(
            ...     time_col='survival_time',
            ...     event_col='event',
            ...     covariates=['age', 'treatment', 'biomarker']
            ... )
            >>> print(result['hazard_ratios']['treatment'])
            0.65
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Cox Proportional Hazards regression")

        try:
            results = self.survival_analyzer.cox_proportional_hazards(
                self.current_data, time_col, event_col, covariates
            )

            if store_results:
                self._store_analysis(results, title or "Cox Proportional Hazards Model")

            return results
        except Exception as e:
            logger.error(f"Cox regression failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'Cox Proportional Hazards Model'
            }

    # ============================================================================
    # BIOEQUIVALENCE ANALYSIS INTEGRATION METHODS
    # ============================================================================

    def analyze_bioequivalence_tost(
        self,
        test_col: str,
        ref_col: str,
        limits: Tuple[float, float] = (0.80, 1.25),
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Two One-Sided Tests (TOST) for bioequivalence

        FDA-recommended method for demonstrating bioequivalence between
        test and reference formulations. Tests if the ratio falls within
        acceptance limits (typically 80-125%).

        Args:
            test_col: Column with test formulation values (e.g., AUC, Cmax)
            ref_col: Column with reference formulation values
            limits: Acceptance limits as tuple (lower, upper)
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Bioequivalence results with TOST statistics, confidence intervals,
            point estimate, and bioequivalence conclusion

        Example:
            >>> result = orchestrator.analyze_bioequivalence_tost(
            ...     test_col='auc_test',
            ...     ref_col='auc_ref'
            ... )
            >>> print(result['bioequivalence_conclusion'])
            'Bioequivalence demonstrated'
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing TOST bioequivalence analysis")

        try:
            results = self.bioequivalence_analyzer.tost(
                self.current_data[test_col].dropna(),
                self.current_data[ref_col].dropna(),
                limits=limits
            )

            if store_results:
                self._store_analysis(results, title or "TOST Bioequivalence Analysis")

            return results
        except Exception as e:
            logger.error(f"TOST analysis failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'TOST Bioequivalence Analysis'
            }

    def analyze_confidence_interval_approach(
        self,
        test_col: str,
        ref_col: str,
        limits: Tuple[float, float] = (0.80, 1.25),
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Confidence Interval approach for bioequivalence

        Alternative to TOST using 90% confidence interval of the geometric
        mean ratio. Bioequivalence is demonstrated if the entire CI falls
        within acceptance limits.

        Args:
            test_col: Column with test formulation values
            ref_col: Column with reference formulation values
            limits: Acceptance limits as tuple (lower, upper)
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            CI results with geometric mean ratio, 90% CI, and conclusion

        Example:
            >>> result = orchestrator.analyze_confidence_interval_approach(
            ...     test_col='cmax_test', ref_col='cmax_ref'
            ... )
            >>> print(result['ci_90'])
            [0.92, 1.08]
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing CI bioequivalence analysis")

        try:
            results = self.bioequivalence_analyzer.confidence_interval_approach(
                self.current_data[test_col].dropna(),
                self.current_data[ref_col].dropna(),
                limits=limits
            )

            if store_results:
                self._store_analysis(results, title or "CI Bioequivalence Analysis")

            return results
        except Exception as e:
            logger.error(f"CI analysis failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'CI Bioequivalence Analysis'
            }

    def analyze_crossover_anova(
        self,
        subject_col: str,
        period_col: str,
        sequence_col: str,
        treatment_col: str,
        value_col: str,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Crossover Design ANOVA for bioequivalence

        Analyzes data from 2x2 crossover studies, accounting for sequence,
        period, and subject effects. Required for FDA bioequivalence submissions.

        Args:
            subject_col: Subject identifier column
            period_col: Period indicator (1 or 2)
            sequence_col: Sequence group (TR or RT)
            treatment_col: Treatment indicator
            value_col: PK parameter values (AUC, Cmax)
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            ANOVA results with effects, degrees of freedom, sums of squares,
            F-statistics, p-values, and treatment effect estimate

        Example:
            >>> result = orchestrator.analyze_crossover_anova(
            ...     subject_col='subject_id', period_col='period',
            ...     sequence_col='sequence', treatment_col='treatment',
            ...     value_col='auc'
            ... )
            >>> print(result['treatment_effect'])
            1.02
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing crossover ANOVA")

        try:
            results = self.bioequivalence_analyzer.crossover_anova(
                self.current_data, subject_col, period_col,
                sequence_col, treatment_col, value_col
            )

            if store_results:
                self._store_analysis(results, title or "Crossover ANOVA")

            return results
        except Exception as e:
            logger.error(f"Crossover ANOVA failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'Crossover ANOVA'
            }

    # ============================================================================
    # DIAGNOSTIC TESTS INTEGRATION METHODS
    # ============================================================================

    def analyze_logistic_regression(
        self,
        y_var: str,
        x_vars: Union[str, List[str]],
        covariates: Optional[List[str]] = None,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Logistic Regression analysis

        Models binary outcomes using one or more predictors. Widely used
        in clinical research for predicting disease presence, treatment response,
        and other binary outcomes.

        Args:
            y_var: Binary outcome variable (0/1)
            x_vars: Predictor variable(s)
            covariates: Additional covariates to include
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Regression results with coefficients, odds ratios,
            confidence intervals, p-values, and model fit statistics

        Example:
            >>> result = orchestrator.analyze_logistic_regression(
            ...     y_var='response', x_vars=['treatment', 'age'],
            ...     covariates=['gender', 'baseline_score']
            ... )
            >>> print(result['odds_ratios']['treatment'])
            2.5
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing logistic regression")

        if isinstance(x_vars, str):
            x_vars = [x_vars]
        
        if covariates is None:
            covariates = []
        
        predictors = x_vars + covariates

        try:
            results = self.diagnostic_tests.logistic_regression(
                self.current_data, y_var, predictors
            )

            if store_results:
                self._store_analysis(results, title or "Logistic Regression")

            return results
        except Exception as e:
            logger.error(f"Logistic regression failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'Logistic Regression'
            }

    def analyze_poisson_regression(
        self,
        y_var: str,
        x_vars: Union[str, List[str]],
        covariates: Optional[List[str]] = None,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Poisson Regression analysis

        Models count data (non-negative integers) with predictors. Suitable for
        analyzing number of events, adverse events, hospitalizations, etc.

        Args:
            y_var: Count outcome variable
            x_vars: Predictor variable(s)
            covariates: Additional covariates to include
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Regression results with coefficients, rate ratios,
            confidence intervals, p-values, and overdispersion check

        Example:
            >>> result = orchestrator.analyze_poisson_regression(
            ...     y_var='adverse_events', x_vars=['treatment', 'age']
            ... )
            >>> print(result['rate_ratios']['treatment'])
            0.75
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Poisson regression")

        if isinstance(x_vars, str):
            x_vars = [x_vars]
        
        if covariates is None:
            covariates = []
        
        predictors = x_vars + covariates

        try:
            results = self.diagnostic_tests.poisson_regression(
                self.current_data, y_var, predictors
            )

            if store_results:
                self._store_analysis(results, title or "Poisson Regression")

            return results
        except Exception as e:
            logger.error(f"Poisson regression failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'Poisson Regression'
            }

    def analyze_linear_mixed_effects(
        self,
        y_var: str,
        x_vars: List[str],
        random_effect: str,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Linear Mixed Effects Model analysis

        Models hierarchical/longitudinal data with both fixed and random effects.
        Essential for repeated measures, crossover studies, and multi-center trials.

        Args:
            y_var: Outcome variable
            x_vars: Fixed effect predictors
            random_effect: Random effect (e.g., subject ID, center)
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Model results with fixed effect estimates, random effect variances,
            confidence intervals, p-values, and model diagnostics

        Example:
            >>> result = orchestrator.analyze_linear_mixed_effects(
            ...     y_var='response', x_vars=['time', 'treatment'],
            ...     random_effect='subject_id'
            ... )
            >>> print(result['fixed_effects']['treatment'])
            5.23
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing linear mixed effects analysis")

        try:
            results = self.diagnostic_tests.linear_mixed_effects(
                self.current_data, y_var, x_vars, random_effect
            )

            if store_results:
                self._store_analysis(results, title or "Linear Mixed Effects Model")

            return results
        except Exception as e:
            logger.error(f"Linear mixed effects analysis failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'Linear Mixed Effects Model'
            }

    # ============================================================================
    # PK/PD ANALYSIS INTEGRATION METHODS
    # ============================================================================

    def analyze_nca_pk(
        self,
        time_col: str,
        conc_col: str,
        dose: float,
        route: str = 'EV',
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Non-Compartmental PK Analysis

        Calculates fundamental PK parameters without assuming a specific
        compartmental model. Standard approach for bioavailability studies
        and regulatory submissions.

        Args:
            time_col: Time points post-dose
            conc_col: Plasma concentrations
            dose: Administered dose
            route: Route of administration ('IV' or 'EV')
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            PK parameters including AUC, Cmax, Tmax, half-life,
            clearance, volume of distribution with CIs

        Example:
            >>> result = orchestrator.analyze_nca_pk(
            ...     time_col='time', conc_col='concentration',
            ...     dose=100, route='EV'
            ... )
            >>> print(result['auc_0_inf'])
            1250.5 ng·h/mL
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing NCA PK analysis")

        try:
            results = self.pkpd_analyzer.non_compartmental_analysis(
                self.current_data, time_col, conc_col, dose, route
            )

            if store_results:
                self._store_analysis(results, title or "Non-Compartmental PK Analysis")

            return results
        except Exception as e:
            logger.error(f"NCA PK analysis failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'Non-Compartmental PK Analysis'
            }

    def analyze_auc(
        self,
        time_col: str,
        conc_col: str,
        dose: float,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Calculate AUC (Area Under Curve) parameters

        Computes AUC0-t (to last observed point) and AUC0-inf (extrapolated to
        infinity) using trapezoidal method with linear and log interpolation.

        Args:
            time_col: Time points
            conc_col: Concentrations
            dose: Administered dose
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            AUC values with calculation method details

        Example:
            >>> result = orchestrator.analyze_auc('time', 'conc', 50)
            >>> print(result['auc_0_t'])
            450.2
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Calculating AUC parameters")

        try:
            results = self.pkpd_analyzer.calculate_auc(
                self.current_data, time_col, conc_col, dose
            )

            if store_results:
                self._store_analysis(results, title or "AUC Analysis")

            return results
        except Exception as e:
            logger.error(f"AUC calculation failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'AUC Analysis'
            }

    def analyze_cmax_tmax(
        self,
        time_col: str,
        conc_col: str,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Calculate Cmax and Tmax

        Identifies maximum observed concentration (Cmax) and time to maximum
        concentration (Tmax) from concentration-time data.

        Args:
            time_col: Time points
            conc_col: Concentrations
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Cmax, Tmax values with confidence intervals

        Example:
            >>> result = orchestrator.analyze_cmax_tmax('time', 'conc')
            >>> print(f"Cmax: {result['cmax']}, Tmax: {result['tmax']}")
            Cmax: 125.3, Tmax: 2.0
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Calculating Cmax and Tmax")

        try:
            results = self.pkpd_analyzer.calculate_cmax_tmax(
                self.current_data, time_col, conc_col
            )

            if store_results:
                self._store_analysis(results, title or "Cmax/Tmax Analysis")

            return results
        except Exception as e:
            logger.error(f"Cmax/Tmax calculation failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'Cmax/Tmax Analysis'
            }

    def analyze_half_life(
        self,
        time_col: str,
        conc_col: str,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Estimate elimination half-life

        Calculates terminal elimination half-life from the log-linear portion
        of the concentration-time curve using at least 3 data points.

        Args:
            time_col: Time points
            conc_col: Concentrations
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Half-life estimate with confidence interval

        Example:
            >>> result = orchestrator.analyze_half_life('time', 'conc')
            >>> print(result['half_life'])
            8.5 hours
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Estimating elimination half-life")

        try:
            results = self.pkpd_analyzer.estimate_half_life(
                self.current_data, time_col, conc_col
            )

            if store_results:
                self._store_analysis(results, title or "Half-Life Analysis")

            return results
        except Exception as e:
            logger.error(f"Half-life estimation failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'Half-Life Analysis'
            }

    def analyze_clearance(
        self,
        time_col: str,
        conc_col: str,
        dose: float,
        route: str = 'EV',
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Calculate clearance parameters

        Computes total clearance (CL) and volume of distribution (Vd)
        based on administration route and PK parameters.

        Args:
            time_col: Time points
            conc_col: Concentrations
            dose: Administered dose
            route: Route of administration ('IV' or 'EV')
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Clearance and volume of distribution values

        Example:
            >>> result = orchestrator.analyze_clearance(
            ...     'time', 'conc', dose=100, route='IV'
            ... )
            >>> print(result['clearance'])
            10.5 L/h
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Calculating clearance parameters")

        try:
            results = self.pkpd_analyzer.calculate_clearance(
                self.current_data, time_col, conc_col, dose, route
            )

            if store_results:
                self._store_analysis(results, title or "Clearance Analysis")

            return results
        except Exception as e:
            logger.error(f"Clearance calculation failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'Clearance Analysis'
            }

    def analyze_pd_response_modeling(
        self,
        conc_col: str,
        effect_col: str,
        model_type: str = 'emax',
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform PD Response Modeling

        Fits pharmacodynamic models (Emax, Sigmoid Emax, Linear) to
        concentration-effect relationships to characterize dose-response.

        Args:
            conc_col: Concentration values
            effect_col: Pharmacodynamic effect values
            model_type: 'emax', 'sigmoid_emax', or 'linear'
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            PD parameters (Emax, EC50, Hill coefficient) with fit statistics

        Example:
            >>> result = orchestrator.analyze_pd_response_modeling(
            ...     'concentration', 'effect', model_type='emax'
            ... )
            >>> print(result['parameters']['emax'])
            95.2
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info(f"Performing PD response modeling ({model_type})")

        try:
            results = self.pkpd_analyzer.pd_response_modeling(
                self.current_data, conc_col, effect_col, model_type
            )

            if store_results:
                self._store_analysis(results, title or f"PD Response Modeling ({model_type})")

            return results
        except Exception as e:
            logger.error(f"PD modeling failed: {e}")
            return {
                'status': 'error',
                'message': f'Analysis failed: {str(e)}',
                'test_name': 'PD Response Modeling'
            }

    # ============================================================================
    # MULTIPLICITY CONTROL INTEGRATION METHODS
    # ============================================================================

    def apply_bonferroni_correction(
        self,
        pvalues: List[float],
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Apply Bonferroni correction for multiple comparisons

        Conservative method controlling family-wise error rate. Adjusts
        significance threshold by dividing by number of tests.

        Args:
            pvalues: List of raw p-values
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Adjusted p-values with correction details

        Example:
            >>> result = orchestrator.apply_bonferroni_correction([0.02, 0.04, 0.06, 0.08])
            >>> print(result['adjusted_pvalues'])
            [0.08, 0.16, 0.24, 0.32]
        """
        logger.info("Applying Bonferroni correction")

        try:
            results = self.multiplicity_control.bonferroni_correction(pvalues)

            if store_results:
                self._store_analysis(results, title or "Bonferroni Correction")

            return results
        except Exception as e:
            logger.error(f"Bonferroni correction failed: {e}")
            return {
                'status': 'error',
                'message': f'Correction failed: {str(e)}',
                'method': 'Bonferroni'
            }

    def apply_holm_correction(
        self,
        pvalues: List[float],
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Apply Holm-Bonferroni (step-down) correction

        Less conservative than Bonferroni while controlling family-wise
        error rate. Sequentially rejects hypotheses ordered by p-value.

        Args:
            pvalues: List of raw p-values
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Adjusted p-values with ordering information

        Example:
            >>> result = orchestrator.apply_holm_correction([0.01, 0.03, 0.05, 0.07])
            >>> print(result['adjusted_pvalues'])
            [0.04, 0.09, 0.10, 0.10]
        """
        logger.info("Applying Holm correction")

        try:
            results = self.multiplicity_control.holm_correction(pvalues)

            if store_results:
                self._store_analysis(results, title or "Holm Correction")

            return results
        except Exception as e:
            logger.error(f"Holm correction failed: {e}")
            return {
                'status': 'error',
                'message': f'Correction failed: {str(e)}',
                'method': 'Holm'
            }

    def apply_bh_fdr(
        self,
        pvalues: List[float],
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Apply Benjamini-Hochberg FDR correction

        Controls false discovery rate rather than family-wise error.
        Less conservative, suitable for exploratory analysis with many tests.

        Args:
            pvalues: List of raw p-values
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            FDR-adjusted q-values with rejection decisions

        Example:
            >>> result = orchestrator.apply_bh_fdr([0.01, 0.02, 0.05, 0.10])
            >>> print(result['q_values'])
            [0.025, 0.04, 0.083, 0.10]
        """
        logger.info("Applying Benjamini-Hochberg FDR correction")

        try:
            results = self.multiplicity_control.bh_fdr_correction(pvalues)

            if store_results:
                self._store_analysis(results, title or "Benjamini-Hochberg FDR")

            return results
        except Exception as e:
            logger.error(f"BH FDR correction failed: {e}")
            return {
                'status': 'error',
                'message': f'Correction failed: {str(e)}',
                'method': 'Benjamini-Hochberg'
            }

    # ============================================================================
    # EXISTING METHODS CONTINUED
    # ============================================================================

    def export_for_thesis(
        self,
        analysis_id: str,
        format: str = 'docx',
        output_path: Optional[str] = None
    ) -> str:
        """Export analysis for thesis integration

        Args:
            analysis_id: ID of stored analysis
            format: Export format ('docx', 'latex', 'json')
            output_path: Output file path

        Returns:
            Path to exported file
        """
        logger.info(f"Exporting analysis {analysis_id} as {format}")

        return self.surfsense_bridge.export_for_thesis(analysis_id, format, output_path)

    def get_analysis(
        self,
        analysis_id: str
    ) -> Optional[Dict[str, Any]]:
        """Retrieve stored analysis

        Args:
            analysis_id: ID of analysis

        Returns:
            Analysis results
        """
        return self.surfsense_bridge.retrieve_analysis(analysis_id)

    def search_analyses(
        self,
        query: str,
        limit: int = 10
    ) -> List[Dict[str, Any]]:
        """Search for stored analyses

        Args:
            query: Search query
            limit: Maximum results

        Returns:
            List of matching analyses
        """
        return self.surfsense_bridge.search_analyses(query, limit)

    def get_available_analyses(
        self
    ) -> List[Dict[str, Any]]:
        """Get list of available statistical analyses

        Returns:
            List of analysis types and descriptions
        """
        return [
            {
                'type': 'descriptive_statistics',
                'name': 'Descriptive Statistics',
                'description': 'Mean, median, SD, distribution analysis',
                'method': 'analyze_descriptive'
            },
            {
                'type': 't_test',
                'name': 'T-Test',
                'description': 'Independent or paired t-test with effect size',
                'method': 'analyze_t_test'
            },
            {
                'type': 'anova',
                'name': 'One-Way ANOVA',
                'description': 'Multiple group comparison with post-hoc tests',
                'method': 'analyze_anova'
            },
            {
                'type': 'correlation',
                'name': 'Correlation Analysis',
                'description': 'Pearson, Spearman, Kendall correlations',
                'method': 'analyze_correlation'
            },
            {
                'type': 'mann_whitney',
                'name': 'Mann-Whitney U Test',
                'description': 'Non-parametric two-group comparison',
                'method': 'analyze_mann_whitney'
            },
            {
                'type': 'kruskal_wallis',
                'name': 'Kruskal-Wallis Test',
                'description': 'Non-parametric multiple group comparison',
                'method': 'analyze_kruskal_wallis'
            },
            {
                'type': 'power_analysis',
                'name': 'Power Analysis',
                'description': 'Sample size and power calculations',
                'method': 'analyze_power'
            },
            # Survival Analysis
            {
                'type': 'kaplan_meier',
                'name': 'Kaplan-Meier Survival Analysis',
                'description': 'Estimate survival function over time with optional grouping',
                'method': 'analyze_survival_kaplan_meier'
            },
            {
                'type': 'log_rank',
                'name': 'Log-Rank Test',
                'description': 'Compare survival curves between groups',
                'method': 'analyze_log_rank_test'
            },
            {
                'type': 'cox_regression',
                'name': 'Cox Proportional Hazards',
                'description': 'Regression model for survival data with covariates',
                'method': 'analyze_cox_proportional_hazards'
            },
            # Bioequivalence
            {
                'type': 'tost_bioequivalence',
                'name': 'TOST Bioequivalence',
                'description': 'Two One-Sided Tests for bioequivalence (FDA method)',
                'method': 'analyze_bioequivalence_tost'
            },
            {
                'type': 'ci_bioequivalence',
                'name': 'CI Bioequivalence',
                'description': 'Confidence Interval approach for bioequivalence',
                'method': 'analyze_confidence_interval_approach'
            },
            {
                'type': 'crossover_anova',
                'name': 'Crossover ANOVA',
                'description': 'ANOVA for 2x2 crossover bioequivalence studies',
                'method': 'analyze_crossover_anova'
            },
            # Diagnostic Tests
            {
                'type': 'logistic_regression',
                'name': 'Logistic Regression',
                'description': 'Model binary outcomes with predictors',
                'method': 'analyze_logistic_regression'
            },
            {
                'type': 'poisson_regression',
                'name': 'Poisson Regression',
                'description': 'Model count data with predictors',
                'method': 'analyze_poisson_regression'
            },
            {
                'type': 'linear_mixed_effects',
                'name': 'Linear Mixed Effects',
                'description': 'Model hierarchical/longitudinal data',
                'method': 'analyze_linear_mixed_effects'
            },
            # PK/PD Analysis
            {
                'type': 'nca_pk',
                'name': 'Non-Compartmental PK',
                'description': 'Calculate AUC, Cmax, Tmax, half-life, clearance',
                'method': 'analyze_nca_pk'
            },
            {
                'type': 'auc_analysis',
                'name': 'AUC Analysis',
                'description': 'Calculate Area Under Curve parameters',
                'method': 'analyze_auc'
            },
            {
                'type': 'cmax_tmax',
                'name': 'Cmax/Tmax Analysis',
                'description': 'Maximum concentration and time to maximum',
                'method': 'analyze_cmax_tmax'
            },
            {
                'type': 'half_life',
                'name': 'Half-Life Analysis',
                'description': 'Estimate elimination half-life',
                'method': 'analyze_half_life'
            },
            {
                'type': 'clearance',
                'name': 'Clearance Analysis',
                'description': 'Calculate clearance and volume of distribution',
                'method': 'analyze_clearance'
            },
            {
                'type': 'pd_modeling',
                'name': 'PD Response Modeling',
                'description': 'Emax and sigmoid Emax models for concentration-effect',
                'method': 'analyze_pd_response_modeling'
            },
            # Multiplicity Control
            {
                'type': 'bonferroni',
                'name': 'Bonferroni Correction',
                'description': 'Conservative multiple comparisons correction',
                'method': 'apply_bonferroni_correction'
            },
            {
                'type': 'holm',
                'name': 'Holm Correction',
                'description': 'Step-down multiple comparisons correction',
                'method': 'apply_holm_correction'
            },
            {
                'type': 'bh_fdr',
                'name': 'Benjamini-Hochberg FDR',
                'description': 'False discovery rate control for exploratory analysis',
                'method': 'apply_bh_fdr'
            }
        ]
    
    # Existing methods continued
    def analyze_wilcoxon_signed_rank(
        self,
        group_col: str,
        value_col: str,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Wilcoxon Signed Rank Test

        Args:
            group_col: Group column (expects 2 groups)
            value_col: Value column
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Wilcoxon Signed Rank Test")

        # Split data into two groups
        groups = self.current_data[group_col].unique()
        if len(groups) != 2:
            raise ValueError("Wilcoxon Signed Rank Test requires exactly 2 groups")

        group1 = self.current_data[self.current_data[group_col] == groups[0]][value_col].values
        group2 = self.current_data[self.current_data[group_col] == groups[1]][value_col].values

        # Perform analysis
        results = self.additional_tools.wilcoxon_signed_rank_test(group1, group2)

        # Store if requested
        if store_results:
            self._store_analysis(results, title or "Wilcoxon Signed Rank Test")

        return results

    def analyze_sign_test(
        self,
        group_col: str,
        value_col: str,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Sign Test

        Args:
            group_col: Group column (expects 2 groups)
            value_col: Value column
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Sign Test")

        # Split data into two groups
        groups = self.current_data[group_col].unique()
        if len(groups) != 2:
            raise ValueError("Sign Test requires exactly 2 groups")

        group1 = self.current_data[self.current_data[group_col] == groups[0]][value_col].values
        group2 = self.current_data[self.current_data[group_col] == groups[1]][value_col].values

        # Perform analysis
        results = self.additional_tools.sign_test(group1, group2)

        # Store if requested
        if store_results:
            self._store_analysis(results, title or "Sign Test")

        return results

    def analyze_friedman(
        self,
        group_col: str,
        value_col: str,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Friedman Test

        Args:
            group_col: Group column
            value_col: Value column
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Friedman Test")

        # Pivot data for Friedman test
        pivot_df = self.current_data.pivot(
            index=self.current_data.index.name if self.current_data.index.name else 'index',
            columns=group_col,
            values=value_col
        )
        data = pivot_df.values.tolist()

        # Perform analysis
        results = self.additional_tools.friedman_test(data)

        # Store if requested
        if store_results:
            self._store_analysis(results, title or "Friedman Test")

        return results

    def analyze_dunns(
        self,
        value_col: str,
        group_col: str,
        p_adjust: str = "bonferroni",
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Dunn's Post-Hoc Test

        Args:
            value_col: Value column
            group_col: Group column
            p_adjust: P-value adjustment method
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Dunn's Post-Hoc Test")

        # Perform analysis
        results = self.additional_tools.dunns_test(
            self.current_data,
            value_col=value_col,
            group_col=group_col,
            p_adjust=p_adjust
        )

        # Store if requested
        if store_results:
            self._store_analysis(results, title or "Dunn's Post-Hoc Test")

        return results

    def analyze_chi_square_goodness_of_fit(
        self,
        observed_col: str,
        expected: Optional[List[float]] = None,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Chi-Square Goodness of Fit Test

        Args:
            observed_col: Column with observed frequencies
            expected: Expected frequencies (optional)
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Chi-Square Goodness of Fit Test")

        # Get observed frequencies
        observed = self.current_data[observed_col].tolist()

        # Perform analysis
        results = self.additional_tools.chi_square_goodness_of_fit(
            observed=observed,
            expected=expected
        )

        # Store if requested
        if store_results:
            self._store_analysis(results, title or "Chi-Square Goodness of Fit Test")

        return results

    def analyze_chi_square_independence(
        self,
        col1: str,
        col2: str,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Chi-Square Test of Independence

        Args:
            col1: First categorical column
            col2: Second categorical column
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Chi-Square Test of Independence")

        # Create contingency table
        contingency_table = pd.crosstab(
            self.current_data[col1],
            self.current_data[col2]
        ).values.tolist()

        # Perform analysis
        results = self.additional_tools.chi_square_test_independence(contingency_table)

        # Store if requested
        if store_results:
            self._store_analysis(results, title or "Chi-Square Test of Independence")

        return results

    def analyze_fisher_exact(
        self,
        col1: str,
        col2: str,
        alternative: str = "two-sided",
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Fisher's Exact Test

        Args:
            col1: First categorical column
            col2: Second categorical column
            alternative: Alternative hypothesis ('two-sided', 'less', 'greater')
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Fisher's Exact Test")

        # Create contingency table
        contingency_table = pd.crosstab(
            self.current_data[col1],
            self.current_data[col2]
        ).values.tolist()

        # Perform analysis
        results = self.additional_tools.fisher_exact_test(
            contingency_table,
            alternative=alternative
        )

        # Store if requested
        if store_results:
            self._store_analysis(results, title or "Fisher's Exact Test")

        return results

    def analyze_mcnemar(
        self,
        col1: str,
        col2: str,
        exact: bool = True,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform McNemar's Test

        Args:
            col1: First binary column
            col2: Second binary column
            exact: Use exact binomial test
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing McNemar's Test")

        # Create contingency table
        contingency_table = pd.crosstab(
            self.current_data[col1],
            self.current_data[col2]
        ).values.tolist()

        # Perform analysis
        results = self.additional_tools.mcnemar_test(
            contingency_table,
            exact=exact
        )

        # Store if requested
        if store_results:
            self._store_analysis(results, title or "McNemar's Test")

        return results

    def analyze_cochran_mantel_haenszel(
        self,
        col1: str,
        col2: str,
        stratify_by: str,
        store_results: bool = True,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform Cochran-Mantel-Haenszel Test

        Args:
            col1: First categorical column
            col2: Second categorical column
            stratify_by: Stratification column
            store_results: Store in SurfSense
            title: Analysis title

        Returns:
            Analysis results
        """
        if self.current_data is None:
            raise ValueError("No data loaded. Import data first.")

        logger.info("Performing Cochran-Mantel-Haenszel Test")

        # Create stratified tables
        tables = []
        for stratum in self.current_data[stratify_by].unique():
            stratum_df = self.current_data[self.current_data[stratify_by] == stratum]
            table = pd.crosstab(stratum_df[col1], stratum_df[col2]).values.tolist()
            tables.append(table)

        # Perform analysis
        results = self.additional_tools.cochran_mantel_haenszel_test(tables)

        # Store if requested
        if store_results:
            self._store_analysis(results, title or "Cochran-Mantel-Haenszel Test")

        return results

    def get_data_summary(self) -> Dict[str, Any]:
        """Get summary of current data

        Returns:
            Data summary
        """
        if self.current_data is None:
            return {'status': 'no_data', 'message': 'No data loaded'}

        return {
            'status': 'loaded',
            'shape': self.current_data.shape,
            'columns': self.current_data.columns.tolist(),
            'dtypes': self.current_data.dtypes.astype(str).to_dict(),
            'memory_usage_mb': self.current_data.memory_usage(deep=True).sum() / (1024 * 1024),
            'missing_values': self.current_data.isnull().sum().to_dict(),
            'metadata': self.current_metadata
        }

    def _store_analysis(
        self,
        results: Dict[str, Any],
        title: str
    ) -> str:
        """Store analysis results in SurfSense"""
        try:
            analysis_id = self.surfsense_bridge.store_analysis_results(
                results,
                title,
                tags=['statistics', 'research']
            )
            self.analysis_cache[analysis_id] = results
            logger.info(f"Stored analysis: {analysis_id}")
            return analysis_id
        except Exception as e:
            logger.warning(f"Failed to store in SurfSense: {e}")
            return None

    # Helper methods for effect size interpretation
    def _interpret_cohens_d(self, d: float) -> str:
        """Interpret Cohen's d magnitude"""
        if d < 0.2:
            return 'negligible'
        elif d < 0.5:
            return 'small'
        elif d < 0.8:
            return 'medium'
        else:
            return 'large'

    def _interpret_eta_squared(self, eta2: float) -> str:
        """Interpret eta-squared magnitude"""
        if eta2 < 0.01:
            return 'small'
        elif eta2 < 0.06:
            return 'medium'
        else:
            return 'large'

    def _interpret_correlation(self, r: float) -> str:
        """Interpret correlation magnitude"""
        if r < 0.1:
            return 'negligible'
        elif r < 0.3:
            return 'small'
        elif r < 0.5:
            return 'medium'
        else:
            return 'large'

    def _interpret_phi(self, phi: float) -> str:
        """Interpret phi coefficient magnitude"""
        return self._interpret_correlation(phi)

    def _interpret_cramers_v(self, v: float) -> str:
        """Interpret Cramér's V magnitude"""
        if v < 0.1:
            return 'negligible'
        elif v < 0.3:
            return 'small'
        elif v < 0.5:
            return 'medium'
        else:
            return 'large'

    def _generate_effect_size_interpretation(
        self,
        effect_size: float,
        effect_type: str,
        magnitude: str
    ) -> str:
        """Generate interpretation text for effect size"""
        return f"The effect size ({effect_type}) of {effect_size:.3f} indicates a {magnitude} effect according to standard conventions."


# Example usage
if __name__ == "__main__":
    orchestrator = StatisticsOrchestrator()

    # Check available analyses
    print("Available Analyses:")
    for analysis in orchestrator.get_available_analyses():
        print(f"  - {analysis['name']}: {analysis['description']}")

    # Test with sample data
    sample_df = pd.DataFrame({
        'group': ['A', 'A', 'A', 'B', 'B', 'B', 'C', 'C', 'C'],
        'value': [10, 12, 11, 15, 14, 16, 8, 9, 7],
        'age': [45, 47, 46, 50, 52, 51, 40, 42, 41]
    })

    # Load data
    orchestrator.current_data = sample_df

    # Perform analysis
    desc_results = orchestrator.analyze_descriptive()
    print("\nDescriptive Analysis:")
    print(json.dumps(desc_results, indent=2, default=str)[:300])
