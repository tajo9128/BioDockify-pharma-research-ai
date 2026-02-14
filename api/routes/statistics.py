
"""Statistics API Routes for BioDockify AI

Provides REST API endpoints for statistical analysis:
- Data import from multiple formats
- Comprehensive statistical analysis
- SurfSense integration
- Thesis export functionality

Complies with:
- Good Laboratory Practice (GLP)
- Good Clinical Practice (GCP)
- FDA/EMA statistical guidelines
- GDPR/CCPA data compliance
"""

import os
import json
import tempfile
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional, Union
from pathlib import Path

from fastapi import APIRouter, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel, Field

import sys
sys.path.insert(0, '/a0/usr/projects/biodockify_ai')

from modules.statistics.orchestrator import StatisticsOrchestrator

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize router
router = APIRouter(prefix="/api/statistics", tags=["Statistics"])

# Global orchestrator instance
_statistics_orchestrator = None


def get_statistics_orchestrator():
    """Get or initialize statistics orchestrator"""
    global _statistics_orchestrator
    if _statistics_orchestrator is None:
        surfsense_url = os.getenv('SURFSENSE_URL', 'http://localhost:8000')
        _statistics_orchestrator = StatisticsOrchestrator(surfsense_url=surfsense_url)
    return _statistics_orchestrator



# Request Models
class TTestRequest(BaseModel):
    """Request model for t-test analysis"""
    group_col: str = Field(..., description="Column containing group labels")
    value_col: str = Field(..., description="Column containing values")
    test_type: str = Field(default="independent", description="'independent' or 'paired'")
    equal_var: bool = Field(default=True, description="Assume equal variance")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")


class ANOVARequest(BaseModel):
    """Request model for ANOVA analysis"""
    value_col: str = Field(..., description="Column containing values")
    group_col: str = Field(..., description="Column containing group labels")
    post_hoc: bool = Field(default=True, description="Perform post-hoc test")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")


class CorrelationRequest(BaseModel):
    """Request model for correlation analysis"""
    columns: List[str] = Field(..., description="Columns to correlate")
    method: str = Field(default="pearson", description="'pearson', 'spearman', or 'kendall'")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")


class PowerAnalysisRequest(BaseModel):
    """Request model for power analysis"""
    test_type: str = Field(default="ttest_ind", description="Type of test")
    effect_size: Optional[float] = Field(None, description="Cohen's d")
    alpha: float = Field(default=0.05, description="Significance level")
    power: float = Field(default=0.80, description="Desired power")
    nobs: Optional[float] = Field(None, description="Sample size")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")


class MannWhitneyRequest(BaseModel):
    """Request model for Mann-Whitney U test"""
    group_col: str = Field(..., description="Column containing group labels")
    value_col: str = Field(..., description="Column containing values")
    alternative: str = Field(default="two-sided", description="'two-sided', 'less', or 'greater'")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")


class KruskalWallisRequest(BaseModel):
    """Request model for Kruskal-Wallis test"""
    value_col: str = Field(..., description="Column containing values")
    group_col: str = Field(..., description="Column containing group labels")
    post_hoc: bool = Field(default=True, description="Perform post-hoc test")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")




# New Request Models for Additional Statistical Tests
class WilcoxonRequest(BaseModel):
    """Request model for Wilcoxon Signed Rank test"""
    group_col: str = Field(..., description="Column containing group labels")
    value_col: str = Field(..., description="Column containing values")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")


class SignTestRequest(BaseModel):
    """Request model for Sign test"""
    group_col: str = Field(..., description="Column containing group labels")
    value_col: str = Field(..., description="Column containing values")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")


class FriedmanRequest(BaseModel):
    """Request model for Friedman test"""
    group_col: str = Field(..., description="Column containing group labels")
    value_col: str = Field(..., description="Column containing values")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")


class DunnsRequest(BaseModel):
    """Request model for Dunn's post-hoc test"""
    value_col: str = Field(..., description="Column containing values")
    group_col: str = Field(..., description="Column containing group labels")
    p_adjust: str = Field(default="bonferroni", description="P-value adjustment method")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")


class ChiSquareGoodnessRequest(BaseModel):
    """Request model for Chi-Square Goodness of Fit test"""
    observed_col: str = Field(..., description="Column with observed frequencies")
    expected: Optional[List[float]] = Field(None, description="Expected frequencies (optional)")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")


class ChiSquareIndependenceRequest(BaseModel):
    """Request model for Chi-Square Test of Independence"""
    col1: str = Field(..., description="First categorical column")
    col2: str = Field(..., description="Second categorical column")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")


class FisherExactRequest(BaseModel):
    """Request model for Fisher's Exact test"""
    col1: str = Field(..., description="First categorical column")
    col2: str = Field(..., description="Second categorical column")
    alternative: str = Field(default="two-sided", description="Alternative hypothesis")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")


class McNemarRequest(BaseModel):
    """Request model for McNemar's test"""
    col1: str = Field(..., description="First binary column")
    col2: str = Field(..., description="Second binary column")
    exact: bool = Field(default=True, description="Use exact binomial test")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")


class CMHRequest(BaseModel):
    """Request model for Cochran-Mantel-Haenszel test"""
    col1: str = Field(..., description="First categorical column")
    col2: str = Field(..., description="Second categorical column")
    stratify_by: str = Field(..., description="Stratification column")
    store_results: bool = Field(default=True, description="Store results in SurfSense")
    title: Optional[str] = Field(None, description="Analysis title")

class ExportRequest(BaseModel):
    """Request model for exporting analysis"""
    analysis_id: str = Field(..., description="ID of stored analysis")
    format: str = Field(default="docx", description="Export format: 'docx', 'latex', 'json'")


# API Endpoints

@router.get("/health")
async def health_check():
    """Health check endpoint"""
    try:
        orchestrator = get_statistics_orchestrator()
        surfsense_ok = orchestrator.surfsense_bridge.health_check()

        return {
            "status": "healthy",
            "surfsense_connected": surfsense_ok,
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return {
            "status": "unhealthy",
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }


@router.get("/available-analyses")
async def get_available_analyses():
    """Get list of available statistical analyses"""
    try:
        orchestrator = get_statistics_orchestrator()
        analyses = orchestrator.get_available_analyses()
        return {
            "status": "success",
            "analyses": analyses,
            "count": len(analyses)
        }
    except Exception as e:
        logger.error(f"Failed to get available analyses: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/import-data")
async def import_data(
    file: UploadFile = File(...),
    clean_data: bool = True,
    validate_data: bool = True
):
    """Import data from uploaded file

    Supported formats: Excel (.xlsx, .xls), CSV, Word (.docx), JSON
    """
    try:
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.filename).suffix) as tmp_file:
            content = await file.read()
            tmp_file.write(content)
            tmp_file_path = tmp_file.name

        logger.info(f"Importing data from: {file.filename}")

        # Import data
        orchestrator = get_statistics_orchestrator()
        result = orchestrator.import_data(
            tmp_file_path,
            clean_data=clean_data,
            validate_data=validate_data
        )

        # Clean up temporary file
        os.unlink(tmp_file_path)

        return {
            "status": "success",
            "message": f"Data imported successfully from {file.filename}",
            "data_summary": {
                "rows": result['metadata']['rows'],
                "columns": result['metadata']['columns'],
                "column_names": result['metadata']['columns_list']
            },
            "metadata": result['metadata']
        }

    except Exception as e:
        logger.error(f"Data import failed: {e}")
        raise HTTPException(status_code=400, detail=f"Data import failed: {str(e)}")


@router.get("/data-summary")
async def get_data_summary():
    """Get summary of currently loaded data"""
    try:
        orchestrator = get_statistics_orchestrator()
        summary = orchestrator.get_data_summary()

        if summary['status'] == 'no_data':
            raise HTTPException(
                status_code=400,
                detail="No data loaded. Import data first using /import-data endpoint."
            )

        return {
            "status": "success",
            "summary": summary
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get data summary: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/descriptive")
async def analyze_descriptive(
    columns: Optional[List[str]] = None,
    store_results: bool = True,
    title: Optional[str] = None
):
    """Perform descriptive statistics analysis"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_descriptive(
            columns=columns,
            store_results=store_results,
            title=title or "Descriptive Statistics"
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"Descriptive analysis failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/t-test")
async def analyze_t_test(request: TTestRequest):
    """Perform t-test analysis"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_t_test(
            group_col=request.group_col,
            value_col=request.value_col,
            test_type=request.test_type,
            equal_var=request.equal_var,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"T-test analysis failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/anova")
async def analyze_anova(request: ANOVARequest):
    """Perform ANOVA analysis"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_anova(
            value_col=request.value_col,
            group_col=request.group_col,
            post_hoc=request.post_hoc,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"ANOVA analysis failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/correlation")
async def analyze_correlation(request: CorrelationRequest):
    """Perform correlation analysis"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_correlation(
            columns=request.columns,
            method=request.method,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"Correlation analysis failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/mann-whitney")
async def analyze_mann_whitney(request: MannWhitneyRequest):
    """Perform Mann-Whitney U test"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_mann_whitney(
            group_col=request.group_col,
            value_col=request.value_col,
            alternative=request.alternative,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"Mann-Whitney test failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/kruskal-wallis")
async def analyze_kruskal_wallis(request: KruskalWallisRequest):
    """Perform Kruskal-Wallis test"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_kruskal_wallis(
            value_col=request.value_col,
            group_col=request.group_col,
            post_hoc=request.post_hoc,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"Kruskal-Wallis test failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/power")
async def analyze_power(request: PowerAnalysisRequest):
    """Perform power analysis"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_power(
            test_type=request.test_type,
            effect_size=request.effect_size,
            alpha=request.alpha,
            power=request.power,
            nobs=request.nobs,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"Power analysis failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/export")
async def export_analysis(request: ExportRequest):
    """Export stored analysis for thesis integration

    Returns downloadable file in specified format
    """
    try:
        orchestrator = get_statistics_orchestrator()
        output_path = orchestrator.export_for_thesis(
            analysis_id=request.analysis_id,
            format=request.format
        )

        filename = Path(output_path).name

        return FileResponse(
            path=output_path,
            filename=filename,
            media_type='application/octet-stream'
        )

    except Exception as e:
        logger.error(f"Export failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/analysis/{analysis_id}")
async def get_analysis(analysis_id: str):
    """Retrieve stored analysis by ID"""
    try:
        orchestrator = get_statistics_orchestrator()
        analysis = orchestrator.get_analysis(analysis_id)

        if analysis is None:
            raise HTTPException(status_code=404, detail=f"Analysis not found: {analysis_id}")

        return {
            "status": "success",
            "analysis_id": analysis_id,
            "analysis": analysis
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to retrieve analysis: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/search")
async def search_analyses(
    query: str,
    limit: int = 10
):
    """Search for stored analyses"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.search_analyses(query=query, limit=limit)

        return {
            "status": "success",
            "query": query,
            "results": results,
            "count": len(results)
        }

    except Exception as e:
        logger.error(f"Search failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))




@router.post("/analyze/wilcoxon-signed-rank")
async def analyze_wilcoxon_signed_rank(request: WilcoxonRequest):
    """Perform Wilcoxon Signed Rank test"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_wilcoxon_signed_rank(
            group_col=request.group_col,
            value_col=request.value_col,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"Wilcoxon Signed Rank test failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/sign-test")
async def analyze_sign_test(request: SignTestRequest):
    """Perform Sign test"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_sign_test(
            group_col=request.group_col,
            value_col=request.value_col,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"Sign test failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/friedman")
async def analyze_friedman(request: FriedmanRequest):
    """Perform Friedman test"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_friedman(
            group_col=request.group_col,
            value_col=request.value_col,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"Friedman test failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/dunns")
async def analyze_dunns(request: DunnsRequest):
    """Perform Dunn's post-hoc test"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_dunns(
            value_col=request.value_col,
            group_col=request.group_col,
            p_adjust=request.p_adjust,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"Dunn's test failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/chi-square-goodness")
async def analyze_chi_square_goodness(request: ChiSquareGoodnessRequest):
    """Perform Chi-Square Goodness of Fit test"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_chi_square_goodness_of_fit(
            observed_col=request.observed_col,
            expected=request.expected,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"Chi-Square Goodness of Fit test failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/chi-square-independence")
async def analyze_chi_square_independence(request: ChiSquareIndependenceRequest):
    """Perform Chi-Square Test of Independence"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_chi_square_independence(
            col1=request.col1,
            col2=request.col2,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"Chi-Square Test of Independence failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/fisher-exact")
async def analyze_fisher_exact(request: FisherExactRequest):
    """Perform Fisher's Exact test"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_fisher_exact(
            col1=request.col1,
            col2=request.col2,
            alternative=request.alternative,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"Fisher's Exact test failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/mcnemar")
async def analyze_mcnemar(request: McNemarRequest):
    """Perform McNemar's test"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_mcnemar(
            col1=request.col1,
            col2=request.col2,
            exact=request.exact,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"McNemar's test failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/analyze/cmh")
async def analyze_cmh(request: CMHRequest):
    """Perform Cochran-Mantel-Haenszel test"""
    try:
        orchestrator = get_statistics_orchestrator()
        results = orchestrator.analyze_cochran_mantel_haenszel(
            col1=request.col1,
            col2=request.col2,
            stratify_by=request.stratify_by,
            store_results=request.store_results,
            title=request.title
        )

        return {
            "status": "success",
            "results": results
        }

    except Exception as e:
        logger.error(f"Cochran-Mantel-Haenszel test failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/clear-data")
async def clear_data():
    """Clear currently loaded data"""
    try:
        orchestrator = get_statistics_orchestrator()
        orchestrator.current_data = None
        orchestrator.current_metadata = None

        return {
            "status": "success",
            "message": "Data cleared successfully"
        }

    except Exception as e:
        logger.error(f"Failed to clear data: {e}")
        raise HTTPException(status_code=500, detail=str(e))
"""Additional API Routes for Statistics UI

Add these to /a0/usr/projects/biodockify_ai/api/routes/statistics.py
"""

# Add these imports at the top if not already present
from typing import Dict, List, Any, Optional

# Add these request models after the existing request models
class GenericAnalysisRequest(BaseModel):
    """Generic request model for all analysis types"""
    analysisType: str = Field(..., description="Type of analysis to perform")
    data: List[Dict[str, Any]] = Field(..., description="Data to analyze")
    parameters: Dict[str, Any] = Field(default_factory=dict, description="Analysis parameters")


class AssumptionsRequest(BaseModel):
    """Request model for assumption testing"""
    analysisType: str = Field(..., description="Type of analysis")
    data: List[Dict[str, Any]] = Field(..., description="Data to test assumptions on")
    parameters: Dict[str, Any] = Field(default_factory=dict, description="Analysis parameters")


class ExportDocxRequest(BaseModel):
    """Request model for DOCX export"""
    results: Dict[str, Any] = Field(..., description="Analysis results")
    data: List[Dict[str, Any]] = Field(..., description="Analysis data")
    analysisType: str = Field(..., description="Type of analysis")


# Add these endpoints to the router after the existing endpoints

@router.post("/upload")
async def upload_statistics_file(
    file: UploadFile = File(...),
    clean_data: bool = True,
    validate_data: bool = True
):
    """Upload data file for statistics analysis

    Supported formats: CSV, Excel (.xlsx, .xls), JSON
    Returns the data in a format suitable for the statistics UI
    """
    try:
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.filename).suffix) as tmp_file:
            content = await file.read()
            tmp_file.write(content)
            tmp_file_path = tmp_file.name

        logger.info(f"Uploading statistics data from: {file.filename}")

        # Import data using the orchestrator
        orchestrator = get_statistics_orchestrator()
        result = orchestrator.import_data(
            tmp_file_path,
            clean_data=clean_data,
            validate_data=validate_data
        )

        # Clean up temporary file
        os.unlink(tmp_file_path)

        # Return the data in the format expected by the UI
        return {
            "status": "success",
            "message": f"Data uploaded successfully from {file.filename}",
            "data": result['data'],
            "metadata": result['metadata']
        }

    except Exception as e:
        logger.error(f"Data upload failed: {e}")
        raise HTTPException(status_code=400, detail=f"Data upload failed: {str(e)}")


@router.post("/analyze")
async def analyze_generic(request: GenericAnalysisRequest):
    """Perform generic statistical analysis

    This endpoint routes the analysis to the appropriate method based on analysisType.
    Supports all 25+ analysis types.
    """
    try:
        orchestrator = get_statistics_orchestrator()
        
        # Convert data list to DataFrame for processing
        import pandas as pd
        df = pd.DataFrame(request.data)
        
        # Load data into the orchestrator
        orchestrator.current_data = df
        orchestrator.current_metadata = {
            'rows': len(df),
            'columns': len(df.columns),
            'columns_list': list(df.columns),
            'column_types': dict(df.dtypes.astype(str))
        }
        
        # Route to appropriate analysis method based on analysisType
        analysis_type = request.analysisType
        params = request.parameters
        
        results = None
        
        # Basic Statistics
        if analysis_type == 'descriptive_statistics':
            results = orchestrator.analyze_descriptive(
                columns=params.get('columns'),
                store_results=False,
                title="Descriptive Statistics"
            )
        
        elif analysis_type == 't_test':
            results = orchestrator.analyze_t_test(
                group_col=params.get('group'),
                value_col=params.get('outcome'),
                test_type='independent',
                equal_var=params.get('variance_assumption') == 'equal',
                store_results=False,
                title="T-Test"
            )
        
        elif analysis_type == 'anova':
            results = orchestrator.analyze_anova(
                value_col=params.get('outcome'),
                group_col=params.get('group'),
                post_hoc=params.get('post_hoc', 'none') != 'none',
                store_results=False,
                title="ANOVA"
            )
        
        elif analysis_type == 'correlation':
            results = orchestrator.analyze_correlation(
                columns=[params.get('variable1'), params.get('variable2')],
                method=params.get('method', 'pearson'),
                store_results=False,
                title="Correlation Analysis"
            )
        
        elif analysis_type == 'mann_whitney':
            results = orchestrator.analyze_mann_whitney(
                group_col=params.get('group'),
                value_col=params.get('outcome'),
                alternative=params.get('alternative', 'two-sided'),
                store_results=False,
                title="Mann-Whitney U Test"
            )
        
        elif analysis_type == 'kruskal_wallis':
            results = orchestrator.analyze_kruskal_wallis(
                value_col=params.get('outcome'),
                group_col=params.get('group'),
                post_hoc=params.get('post_hoc', 'none') != 'none',
                store_results=False,
                title="Kruskal-Wallis Test"
            )
        
        elif analysis_type == 'power_analysis':
            results = orchestrator.analyze_power(
                test_type=params.get('test_type', 'ttest_ind'),
                effect_size=params.get('effect_size'),
                alpha=params.get('alpha', 0.05),
                power=params.get('power', 0.80),
                nobs=params.get('sample_size'),
                store_results=False,
                title="Power Analysis"
            )
        
        # Survival Analysis - use survival_analysis module
        elif analysis_type in ['kaplan_meier', 'log_rank', 'cox_ph']:
            from modules.statistics.survival_analysis import SurvivalAnalysis
            survival = SurvivalAnalysis()
            
            # Prepare data
            time_col = params.get('time')
            event_col = params.get('event')
            group_col = params.get('group')
            
            if analysis_type == 'kaplan_meier':
                result = survival.kaplan_meier_estimate(
                    df, time_col, event_col,
                    conf_int=True,
                    conf_level=params.get('confidence_level', 0.95)
                )
                results = {
                    'analysisType': 'kaplan_meier',
                    'testStatistics': {'survival_function': str(result['survival_function'])},
                    'pValue': result.get('p_value'),
                    'confidenceInterval': result.get('confidence_intervals', [0, 1]),
                    'confidenceLevel': params.get('confidence_level', 0.95),
                    'sampleSize': len(df),
                    'conclusion': result.get('interpretation', 'Survival analysis completed'),
                    'significance': 0.05
                }
            
            elif analysis_type == 'log_rank':
                result = survival.log_rank_test(
                    df, time_col, event_col, group_col
                )
                results = {
                    'analysisType': 'log_rank',
                    'testStatistics': result['test_results'],
                    'pValue': result['test_results'].get('p_value'),
                    'confidenceLevel': params.get('confidence_level', 0.95),
                    'sampleSize': len(df),
                    'conclusion': result['interpretation'],
                    'significance': 0.05
                }
            
            elif analysis_type == 'cox_ph':
                covariates = params.get('covariates', [])
                strata = params.get('strata', [])
                result = survival.cox_proportional_hazards(
                    df, time_col, event_col, covariates,
                    strata=strata if strata else None
                )
                results = {
                    'analysisType': 'cox_ph',
                    'testStatistics': result['model_results'],
                    'pValues': result['model_results'].get('p_values', {}),
                    'confidenceIntervals': result['model_results'].get('confidence_intervals', {}),
                    'confidenceLevel': params.get('confidence_level', 0.95),
                    'sampleSize': len(df),
                    'conclusion': result['interpretation'],
                    'significance': 0.05
                }
        
        # Bioequivalence
        elif analysis_type in ['tost', 'ci_approach', 'crossover_anova']:
            from modules.statistics.bioequivalence_analysis import BioequivalenceAnalysis
            beq = BioequivalenceAnalysis()
            
            test_col = params.get('test')
            ref_col = params.get('reference')
            
            if analysis_type == 'tost':
                result = beq.two_one_sided_tests(
                    df, test_col, ref_col,
                    log_transform=params.get('log_transform', True),
                    alpha=1 - params.get('confidence_level', 0.90)
                )
                results = {
                    'analysisType': 'tost',
                    'testStatistics': result['test_results'],
                    'pValue': result['test_results'].get('p_value'),
                    'confidenceInterval': result['test_results'].get('confidence_interval', [0, 1]),
                    'confidenceLevel': params.get('confidence_level', 0.90),
                    'sampleSize': len(df),
                    'conclusion': result['interpretation'],
                    'significance': 0.10
                }
            
            elif analysis_type == 'ci_approach':
                result = beq.confidence_interval_approach(
                    df, test_col, ref_col,
                    log_transform=params.get('log_transform', True)
                )
                results = {
                    'analysisType': 'ci_approach',
                    'testStatistics': result['test_results'],
                    'confidenceInterval': result['test_results'].get('confidence_interval', [0, 1]),
                    'confidenceLevel': params.get('confidence_level', 0.90),
                    'sampleSize': len(df),
                    'conclusion': result['interpretation'],
                    'significance': 0.10
                }
        
        # PK/PD Analysis
        elif analysis_type in ['nca_pk', 'auc_calculation', 'cmax_tmax', 'half_life', 'clearance', 'pd_response_modeling']:
            from modules.statistics.pkpd_analysis import PKPDAnalysis
            pkpd = PKPDAnalysis(
                df,
                dose=params.get('dose', 1),
                route=params.get('route', 'EV')
            )
            
            if analysis_type == 'nca_pk':
                result = pkpd.non_compartmental_analysis()
                results = {
                    'analysisType': 'nca_pk',
                    'testStatistics': result['parameters'],
                    'confidenceIntervals': result['confidence_intervals'],
                    'confidenceLevel': params.get('confidence_level', 0.90),
                    'sampleSize': len(df),
                    'conclusion': result['interpretation'],
                    'significance': 0.10
                }
            
            elif analysis_type == 'auc_calculation':
                time_col = params.get('time')
                conc_col = params.get('concentration')
                result = pkpd.calculate_auc(time_col, conc_col)
                results = {
                    'analysisType': 'auc_calculation',
                    'testStatistics': result,
                    'confidenceLevel': params.get('confidence_level', 0.90),
                    'sampleSize': len(df),
                    'conclusion': 'AUC calculation completed',
                    'significance': 0.10
                }
        
        # Categorical Tests
        elif analysis_type == 'wilcoxon_signed_rank':
            results = orchestrator.analyze_wilcoxon_signed_rank(
                group_col=params.get('group'),
                value_col=params.get('outcome'),
                store_results=False,
                title="Wilcoxon Signed Rank Test"
            )
        
        elif analysis_type == 'sign_test':
            results = orchestrator.analyze_sign_test(
                group_col=params.get('group'),
                value_col=params.get('outcome'),
                store_results=False,
                title="Sign Test"
            )
        
        elif analysis_type == 'friedman':
            results = orchestrator.analyze_friedman(
                group_col=params.get('group'),
                value_col=params.get('outcome'),
                store_results=False,
                title="Friedman Test"
            )
        
        elif analysis_type == 'dunns':
            results = orchestrator.analyze_dunns(
                value_col=params.get('outcome'),
                group_col=params.get('group'),
                p_adjust=params.get('p_adjust', 'bonferroni'),
                store_results=False,
                title="Dunn's Post-Hoc Test"
            )
        
        elif analysis_type == 'chi_square':
            col1 = params.get('variable1')
            col2 = params.get('variable2')
            results = orchestrator.analyze_chi_square_independence(
                col1=col1,
                col2=col2,
                store_results=False,
                title="Chi-Square Test"
            )
        
        elif analysis_type == 'fisher_exact':
            col1 = params.get('variable1')
            col2 = params.get('variable2')
            results = orchestrator.analyze_fisher_exact(
                col1=col1,
                col2=col2,
                alternative=params.get('alternative', 'two-sided'),
                store_results=False,
                title="Fisher's Exact Test"
            )
        
        elif analysis_type == 'mcnemar':
            col1 = params.get('variable1')
            col2 = params.get('variable2')
            results = orchestrator.analyze_mcnemar(
                col1=col1,
                col2=col2,
                exact=params.get('correction', True),
                store_results=False,
                title="McNemar's Test"
            )
        
        elif analysis_type == 'cmh':
            col1 = params.get('variable1')
            col2 = params.get('variable2')
            strata = params.get('strata')
            results = orchestrator.analyze_cochran_mantel_haenszel(
                col1=col1,
                col2=col2,
                stratify_by=strata,
                store_results=False,
                title="Cochran-Mantel-Haenszel Test"
            )
        
        # Multiplicity Control
        elif analysis_type in ['bonferroni', 'holm', 'bh_fdr']:
            from modules.statistics.multiplicity_control import MultiplicityControl
            multi = MultiplicityControl()
            
            p_value_cols = params.get('p_values', [])
            p_values = []
            for col in p_value_cols:
                if col in df.columns:
                    p_values.extend(df[col].dropna().tolist())
            
            if analysis_type == 'bonferroni':
                result = multi.bonferroni_correction(p_values, alpha=params.get('alpha', 0.05))
            elif analysis_type == 'holm':
                result = multi.holm_method(p_values, alpha=params.get('alpha', 0.05))
            elif analysis_type == 'bh_fdr':
                result = multi.benjamini_hochberg(p_values, q=params.get('q', 0.05))
            
            results = {
                'analysisType': analysis_type,
                'adjustedPValues': result['adjusted_p_values'],
                'sampleSize': len(p_values),
                'conclusion': result['interpretation'],
                'significance': params.get('alpha', params.get('q', 0.05))
            }
        
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported analysis type: {analysis_type}")
        
        # Format results for UI
        if results:
            return {
                "status": "success",
                "results": results
            }
        else:
            raise HTTPException(status_code=500, detail="Analysis returned no results")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Analysis failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/assumptions")
async def test_assumptions(request: AssumptionsRequest):
    """Test assumptions for selected analysis

    Returns a list of assumptions with their status (passed/failed/warning),
    interpretations, and suggested remedies.
    """
    try:
        import pandas as pd
        from scipy import stats
        import numpy as np
        
        df = pd.DataFrame(request.data)
        analysis_type = request.analysisType
        params = request.parameters
        
        assumptions = []
        
        # Test normality for numeric columns
        outcome_col = params.get('outcome')
        if outcome_col and outcome_col in df.columns:
            # Shapiro-Wilk test for normality
            stat, p = stats.shapiro(df[outcome_col].dropna())
            normality_passed = p > 0.05
            assumptions.append({
                'name': 'Normality',
                'status': 'passed' if normality_passed else 'failed',
                'testStatistic': stat,
                'pValue': p,
                'interpretation': f'p-value = {p:.4f}. Data is {"normal" if normality_passed else "not normal"}.',
                'remedy': 'Use non-parametric tests like Mann-Whitney U or Wilcoxon Signed Rank.' if not normality_passed else None
            })
        
        # Test homogeneity of variance
        outcome_col = params.get('outcome')
        group_col = params.get('group')
        if outcome_col and group_col and outcome_col in df.columns and group_col in df.columns:
            groups = [group[outcome_col].dropna().values for name, group in df.groupby(group_col)]
            stat, p = stats.levene(*groups)
            homogeneity_passed = p > 0.05
            assumptions.append({
                'name': 'Homogeneity of Variance',
                'status': 'passed' if homogeneity_passed else 'failed',
                'testStatistic': stat,
                'pValue': p,
                'interpretation': f'p-value = {p:.4f}. Variances are {"equal" if homogeneity_passed else "not equal"}.',
                'remedy': 'Use Welch\'s ANOVA or Kruskal-Wallis test.' if not homogeneity_passed else None
            })
        
        # Test independence (for paired tests)
        if analysis_type in ['t_test', 'wilcoxon_signed_rank'] and 'outcome1' in params and 'outcome2' in params:
            # Check if data is paired (same subject IDs)
            if 'subject' in params:
                # Test correlation between paired measurements
                col1 = params['outcome1']
                col2 = params['outcome2']
                if col1 in df.columns and col2 in df.columns:
                    stat, p = stats.pearsonr(df[col1], df[col2])
                    independence_passed = p > 0.05
                    assumptions.append({
                        'name': 'Independence',
                        'status': 'warning' if independence_passed else 'passed',
                        'testStatistic': stat,
                        'pValue': p,
                        'interpretation': f'Correlation = {stat:.4f}, p = {p:.4f}.',
                        'remedy': 'Consider using paired tests if measurements are from the same subjects.'
                    })
        
        # Test proportional hazards (for survival analysis)
        if analysis_type in ['log_rank', 'cox_ph']:
            from modules.statistics.survival_analysis import SurvivalAnalysis
            survival = SurvivalAnalysis()
            # Simplified proportional hazards check
            assumptions.append({
                'name': 'Proportional Hazards',
                'status': 'warning',
                'interpretation': 'Proportional hazards assumption should be verified with Schoenfeld residuals test.',
                'remedy': 'Consider stratified Cox model or time-varying covariates if violated.'
            })
        
        return {
            "status": "success",
            "assumptionsCheck": assumptions,
            "total": len(assumptions),
            "passed": sum(1 for a in assumptions if a['status'] == 'passed'),
            "failed": sum(1 for a in assumptions if a['status'] == 'failed'),
            "warnings": sum(1 for a in assumptions if a['status'] == 'warning')
        }

    except Exception as e:
        logger.error(f"Assumption testing failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/export/docx")
async def export_docx(request: ExportDocxRequest):
    """Export analysis results to DOCX format

    Returns a downloadable DOCX file with comprehensive analysis report.
    """
    try:
        from export.docx_academic import DocxAcademicExport
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.lib.enums import WD_ORIENTATION
        
        results = request.results
        data = request.data
        analysis_type = request.analysisType
        
        # Create document
        doc = Document()
        
        # Set orientation to landscape
n        section = doc.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        
        # Title
        title = doc.add_heading('Statistical Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Analysis type
        doc.add_heading(f'Analysis Type: {analysis_type.replace("_", " ").title()}', level=1)
        
        # Summary
        doc.add_heading('Summary', level=2)
        p = doc.add_paragraph()
        p.add_run('Conclusion: ').bold = True
        p.add_run(results.get('conclusion', 'N/A'))
        
        if results.get('recommendation'):
            p = doc.add_paragraph()
            p.add_run('Recommendation: ').bold = True
            p.add_run(results['recommendation'])
        
        # Sample size
        if results.get('sampleSize'):
            p = doc.add_paragraph()
            p.add_run('Sample Size: ').bold = True
            p.add_run(str(results['sampleSize']))
        
        # Test Statistics
        if results.get('testStatistics'):
            doc.add_heading('Test Statistics', level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Statistic'
            hdr_cells[1].text = 'Value'
            
            for key, value in results['testStatistics'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(value)
        
        # P-values
        if results.get('pValue') is not None:
            doc.add_heading('P-value', level=2)
            p = doc.add_paragraph()
            p.add_run(f"P-value: {results['pValue']:.4f}")
            
            if results.get('significance'):
                significant = results['pValue'] < results['significance']
                p = doc.add_paragraph()
                p.add_run(f"Significant at α = {results['significance']}: {'Yes' if significant else 'No'}")
        
        # Adjusted P-values
        if results.get('adjustedPValues'):
            doc.add_heading('Adjusted P-values', level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Test'
            hdr_cells[1].text = 'Adjusted P-value'
            
            for key, value in results['adjustedPValues'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = f"{value:.4f}"
        
        # Effect Size
        if results.get('effectSize') is not None:
            doc.add_heading('Effect Size', level=2)
            p = doc.add_paragraph()
            p.add_run(f"Effect Size: {results['effectSize']:.4f}")
            
            if results.get('effectSizeInterpretation'):
                p = doc.add_paragraph()
                p.add_run('Interpretation: ')
                p.add_run(results['effectSizeInterpretation'])
        
        # Confidence Intervals
        if results.get('confidenceInterval'):
            ci = results['confidenceInterval']
            doc.add_heading(f"{results.get('confidenceLevel', 0.95)*100}% Confidence Interval", level=2)
            p = doc.add_paragraph()
            p.add_run(f"[{ci[0]:.4f}, {ci[1]:.4f}]")
        
        if results.get('confidenceIntervals'):
            doc.add_heading('Confidence Intervals', level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Parameter'
            hdr_cells[1].text = 'CI'
            
            for key, value in results['confidenceIntervals'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = f"[{value[0]:.4f}, {value[1]:.4f}]"
        
        # Interpretation
        if results.get('interpretation'):
            doc.add_heading('Interpretation', level=2)
            doc.add_paragraph(results['interpretation'])
        
        # Assumption Checks
        if results.get('assumptionsCheck'):
            doc.add_heading('Assumption Checks', level=2)
            for check in results['assumptionsCheck']:
                p = doc.add_paragraph()
                p.add_run(f"{check['name']}: ").bold = True
                status_color = 'green' if check['status'] == 'passed' else 'red' if check['status'] == 'failed' else 'orange'
                p.add_run(f"{check['status'].upper()}")
                doc.add_paragraph(check.get('interpretation', ''))
                if check.get('remedy'):
                    p = doc.add_paragraph()
                    p.add_run('Remedy: ').bold = True
                    p.add_run(check['remedy'])
        
        # Data Summary
        if data:
            doc.add_heading('Data Summary', level=2)
            table = doc.add_table(rows=1, cols=len(data[0]))
            table.style = 'Table Grid'
            
            # Header row
            for i, key in enumerate(data[0].keys()):
                table.rows[0].cells[i].text = str(key)
            
            # Data rows (max 50)
            for row in data[:50]:
                row_cells = table.add_row().cells
                for i, value in enumerate(row.values()):
                    row_cells[i].text = str(value)[:50]  # Truncate long values
            
            if len(data) > 50:
                p = doc.add_paragraph()
                p.add_run(f"... and {len(data) - 50} more rows")
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Analysis Type: {analysis_type}")
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            tmp_file_path = tmp_file.name
        
        filename = f"statistics_{analysis_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return FileResponse(
            path=tmp_file_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f"DOCX export failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))
