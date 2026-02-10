"""
Academic DOCX Export with Templates

This module provides comprehensive DOCX document generation for academic writing,
including:
- Multiple academic templates (IEEE, APA, Nature, custom)
- Title page generation
- Chapter and section structure
- Figure and table insertion
- Abstract formatting
- Reference list formatting
- Custom styling (fonts, spacing, alignment)
- Page layout control
- Header and footer support

Export Pipeline:
Research Data → Template Selection → Content Assembly → DOCX → Styling → Output Files
"""

import pandas as pd
from pathlib import Path
from typing import List, Dict, Optional, Union, Any
from dataclasses import dataclass
import logging


try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. Install with: pip install python-docx")


# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class FigureData:
    """Data structure for figures in DOCX"""
    path: str
    caption: str
    width_inches: float = 5.0
    position: str = 'center'


@dataclass
class TableData:
    """Data structure for tables in DOCX"""
    headers: List[str]
    rows: List[List[Any]]
    caption: Optional[str] = None
    style: str = 'Light Grid Accent 1'


class AcademicDOCXExporter:
    """
    Academic DOCX document exporter with template support

    Features:
    - Multiple academic templates (IEEE, APA, Nature)
    - Custom styling and formatting
    - Chapter and section structure
    - Figure and table insertion
    - Title page generation
    - Abstract formatting
    - Reference list formatting
    - Page layout control
    - Header and footer support

    Supported Templates:
    - IEEE: Institute of Electrical and Electronics Engineers
    - APA: American Psychological Association
    - Nature: Nature journal style
    - Custom: User-defined styling
    """

    # Template configurations
    TEMPLATES = {
        'ieee': {
            'font_name': 'Times New Roman',
            'font_size': 10,
            'line_spacing': 1.0,
            'margin_top': 1.0,
            'margin_bottom': 1.0,
            'margin_left': 1.0,
            'margin_right': 1.0
        },
        'apa': {
            'font_name': 'Times New Roman',
            'font_size': 12,
            'line_spacing': 2.0,
            'margin_top': 1.0,
            'margin_bottom': 1.0,
            'margin_left': 1.0,
            'margin_right': 1.0
        },
        'nature': {
            'font_name': 'Arial',
            'font_size': 11,
            'line_spacing': 1.5,
            'margin_top': 0.75,
            'margin_bottom': 0.75,
            'margin_left': 0.75,
            'margin_right': 0.75
        }
    }

    def __init__(
        self,
        template: str = 'ieee',
        custom_config: Optional[Dict] = None
    ):
        """
        Initialize academic DOCX exporter

        Args:
            template: Template name ('ieee', 'apa', 'nature', 'custom')
            custom_config: Custom configuration for 'custom' template
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "The 'python-docx' library is required for academic document generation. "
                "Please install it using: pip install python-docx"
            )

        self.doc = Document()
        self.template = template
        self.config = self._get_template_config(template, custom_config)

        # Apply template styling
        self._apply_template()

        logger.info(f"DOCX exporter initialized with template: {template}")

    def _get_template_config(self, template: str, custom_config: Optional[Dict]) -> Dict:
        """Get template configuration"""
        if template == 'custom' and custom_config:
            return custom_config
        return self.TEMPLATES.get(template, self.TEMPLATES['ieee'])

    def _apply_template(self):
        """Apply template styling to document"""

        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.config['font_name']
        font.size = Pt(self.config['font_size'])

        # Set line spacing
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = self.config['line_spacing']
        paragraph_format.space_after = Pt(6)

        # Set page margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(self.config['margin_top'])
            section.bottom_margin = Inches(self.config['margin_bottom'])
            section.left_margin = Inches(self.config['margin_left'])
            section.right_margin = Inches(self.config['margin_right'])

        logger.debug(f"Applied template: {self.template}")

    def _set_ieee_styles(self):
        """IEEE template: Times New Roman, 10pt"""
        self._apply_font_style('Times New Roman', 10, 1.0)

    def _set_apa_styles(self):
        """APA template: Times New Roman, 12pt, double spaced"""
        self._apply_font_style('Times New Roman', 12, 2.0)

    def _set_nature_styles(self):
        """Nature template: Arial, 11pt, 1.5 spaced"""
        self._apply_font_style('Arial', 11, 1.5)

    def _apply_font_style(self, font_name: str, font_size: int, line_spacing: float):
        """Apply font style to default style"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)

        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = line_spacing

    def set_margins(
        self,
        top: float,
        bottom: float,
        left: float,
        right: float
    ):
        """
        Set page margins (in inches)

        Args:
            top: Top margin
            bottom: Bottom margin
            left: Left margin
            right: Right margin
        """
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)

        logger.debug(f"Margins set: T={top}, B={bottom}, L={left}, R={right}")

    def add_title_page(
        self,
        title: str,
        author: str,
        institution: str,
        date: Optional[str] = None
    ):
        """
        Add thesis title page

        Args:
            title: Document title
            author: Author name
            institution: Institution name
            date: Optional date (default: current date)
        """
        # Title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = self.config['font_name']

        self.doc.add_paragraph()  # Spacing

        # Author
        author_para = self.doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_para.add_run(f"by\n{author}").font.size = Pt(12)

        self.doc.add_paragraph()

        # Institution
        inst_para = self.doc.add_paragraph()
        inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inst_para.add_run(institution).font.size = Pt(12)

        # Date
        if date:
            self.doc.add_paragraph()
            date_para = self.doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.add_run(date).font.size = Pt(12)
        else:
            from datetime import datetime
            self.doc.add_paragraph()
            date_para = self.doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.add_run(datetime.now().strftime("%B %Y")).font.size = Pt(12)

        self.doc.add_page_break()

        logger.debug("Title page added")

    def add_abstract(self, text: str, heading: str = 'Abstract'):
        """
        Add formatted abstract

        Args:
            text: Abstract text
            heading: Abstract heading (default: "Abstract")
        """
        # Heading
        self.doc.add_heading(heading, level=1)

        # Abstract text
        para = self.doc.add_paragraph(text)
        para.paragraph_format.first_line_indent = Inches(0.5)

        self.doc.add_page_break()

        logger.debug("Abstract added")

    def add_keywords(self, keywords: List[str]):
        """
        Add keywords section

        Args:
            keywords: List of keywords
        """
        para = self.doc.add_paragraph()
        run = para.add_run("Keywords: ")
        run.font.bold = True

        run = para.add_run(", ".join(keywords))
        run.font.italic = True

        self.doc.add_paragraph()

        logger.debug(f"Added {len(keywords)} keywords")

    def add_chapter(
        self,
        number: int,
        title: str,
        sections: List[Dict],
        add_page_break: bool = False
    ):
        """
        Add chapter with sections

        Args:
            number: Chapter number
            title: Chapter title
            sections: List of section dictionaries
            add_page_break: Whether to add page break after chapter
        """
        # Chapter heading
        chapter_heading = f'Chapter {number}: {title}'
        self.doc.add_heading(chapter_heading, level=1)

        # Add sections
        for section in sections:
            self._add_section(section)

        if add_page_break:
            self.doc.add_page_break()

        logger.debug(f"Added chapter {number}: {title}")

    def _add_section(self, section: Dict, level: int = 2):
        """
        Add section or subsection

        Args:
            section: Section data with title, text, figures, tables
            level: Section level (2=section, 3=subsection)
        """
        # Section heading
        self.doc.add_heading(section.get('title', 'Untitled'), level=level)

        # Section content
        text = section.get('text', '')
        if text:
            self.doc.add_paragraph(text)

        # Add subsections
        for subsection in section.get('subsections', []):
            self._add_section(subsection, level=3)

        # Add tables
        for table_data in section.get('tables', []):
            self.add_table(table_data)

        # Add figures
        for figure in section.get('figures', []):
            self._add_figure(figure)

    def add_section(
        self,
        title: str,
        content: str,
        level: int = 2
    ):
        """
        Add a simple section

        Args:
            title: Section title
            content: Section content text
            level: Heading level (1-3)
        """
        self.doc.add_heading(title, level=level)
        self.doc.add_paragraph(content)

    def add_paragraph(
        self,
        text: str,
        bold: bool = False,
        italic: bool = False,
        align: str = 'left',
        first_line_indent: float = 0.0
    ):
        """
        Add a formatted paragraph

        Args:
            text: Paragraph text
            bold: Whether text is bold
            italic: Whether text is italic
            align: Text alignment ('left', 'center', 'right', 'justify')
            first_line_indent: First line indent in inches
        """
        para = self.doc.add_paragraph()

        # Set alignment
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        para.alignment = alignment_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)

        # Add text
        run = para.add_run(text)
        run.font.bold = bold
        run.font.italic = italic

        # First line indent
        if first_line_indent > 0:
            para.paragraph_format.first_line_indent = Inches(first_line_indent)

        return para

    def add_table(self, data: Union[Dict, TableData]):
        """
        Add formatted table

        Args:
            data: Table data as dict or TableData object with 'headers', 'rows', 'caption'
        """
        if isinstance(data, dict):
            table_data = TableData(
                headers=data['headers'],
                rows=data['rows'],
                caption=data.get('caption'),
                style=data.get('style', 'Light Grid Accent 1')
            )
        else:
            table_data = data

        # Create table
        table = self.doc.add_table(
            rows=1,
            cols=len(table_data.headers)
        )
        table.style = table_data.style

        # Headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(table_data.headers):
            hdr_cells[i].text = header
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.font.bold = True

        # Data rows
        for row_data in table_data.rows:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = str(value)

        # Caption
        if table_data.caption:
            caption = self.doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(f"Table: {table_data.caption}")
            run.font.italic = True

        self.doc.add_paragraph()  # Spacing

        logger.debug("Table added")

    def add_table_from_dataframe(
        self,
        df: pd.DataFrame,
        caption: Optional[str] = None,
        style: str = 'Light Grid Accent 1'
    ):
        """
        Add table from pandas DataFrame

        Args:
            df: Pandas DataFrame
            caption: Optional table caption
            style: Table style
        """
        table_data = TableData(
            headers=list(df.columns),
            rows=[list(row) for _, row in df.iterrows()],
            caption=caption,
            style=style
        )

        self.add_table(table_data)

        logger.debug(f"Table added from DataFrame ({len(df)} rows)")

    def _add_figure(self, figure: Union[Dict, FigureData]):
        """
        Add figure with caption

        Args:
            figure: Figure data as dict or FigureData object
        """
        if isinstance(figure, dict):
            fig_data = FigureData(
                path=figure['path'],
                caption=figure['caption'],
                width_inches=figure.get('width_inches', 5.0)
            )
        else:
            fig_data = figure

        # Add image
        self.doc.add_picture(fig_data.path, width=Inches(fig_data.width_inches))

        # Caption
        caption_para = self.doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_para.add_run(f"Figure: {fig_data.caption}")
        run.font.italic = True

        self.doc.add_paragraph()  # Spacing

        logger.debug(f"Figure added: {fig_data.path}")

    def add_figure(self, image_path: str, caption: str, width_inches: float = 5.0):
        """
        Add figure with caption (simplified interface)

        Args:
            image_path: Path to image file
            caption: Figure caption
            width_inches: Width in inches
        """
        self._add_figure({
            'path': image_path,
            'caption': caption,
            'width_inches': width_inches
        })

    def add_references(self, references: List[str]):
        """
        Add reference list

        Args:
            references: List of reference strings
        """
        self.doc.add_heading('References', level=1)

        for i, ref in enumerate(references, 1):
            ref_para = self.doc.add_paragraph(f"[{i}] {ref}")
            ref_para.paragraph_format.left_indent = Inches(0.5)
            ref_para.paragraph_format.space_after = Pt(6)

        logger.debug(f"Added {len(references)} references")

    def add_bibliography_from_bibtex(
        self,
        bibtex_file: str,
        style: str = 'apa'
    ):
        """
        Add bibliography from BibTeX file

        Args:
            bibtex_file: Path to .bib file
            style: Citation style for formatting
        """
        try:
            from export.bibtex_manager import BibTeXManager

            # Load BibTeX file
            bib_manager = BibTeXManager()
            bib_manager.load_from_file(bibtex_file)

            # Format all citations
            citations = bib_manager.format_all_citations(style)

            # Add to document
            self.add_references(citations)

            logger.debug(f"Added bibliography from {bibtex_file} with {len(citations)} entries")

        except ImportError:
            logger.error("bibtex_manager not available")
            raise

    def add_page_break(self):
        """Add page break"""
        self.doc.add_page_break()

    def add_horizontal_line(self):
        """Add horizontal line"""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('_' * 80)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def add_bullet_list(self, items: List[str]):
        """
        Add bullet point list

        Args:
            items: List of bullet items
        """
        for item in items:
            para = self.doc.add_paragraph(item, style='List Bullet')
            para.paragraph_format.space_after = Pt(3)

        logger.debug(f"Added bullet list with {len(items)} items")

    def add_numbered_list(self, items: List[str]):
        """
        Add numbered list

        Args:
            items: List of numbered items
        """
        for item in items:
            para = self.doc.add_paragraph(item, style='List Number')
            para.paragraph_format.space_after = Pt(3)

        logger.debug(f"Added numbered list with {len(items)} items")

    def add_header(self, text: str):
        """
        Add header to all pages

        Args:
            text: Header text
        """
        section = self.doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        logger.debug(f"Header added: {text}")

    def add_footer(self, text: str, include_page_number: bool = False):
        """
        Add footer to all pages

        Args:
            text: Footer text
            include_page_number: Whether to include page number
        """
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]

        if include_page_number:
            footer_para.text = f"{text} Page "
        else:
            footer_para.text = text

        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if include_page_number:
            # Add page number field
            from docx.oxml.shared import qn
            fldChar1 = self.doc.oxml.shared.OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = self.doc.oxml.shared.OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"

            fldChar2 = self.doc.oxml.shared.OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            footer_para._p.append(fldChar1)
            footer_para._p.append(instrText)
            footer_para._p.append(fldChar2)
        logger.debug(f"Footer added: {text}")

    def export(self, filepath: str) -> str:
        """
        Save DOCX file using atomic temp file replacement (Bug #16).
        
        Args:
            filepath: Path to output .docx file
        """
        import tempfile
        import os
        
        output_path = Path(filepath)

        # Ensure .docx extension
        if not output_path.suffix == '.docx':
            output_path = output_path.with_suffix('.docx')

        # Create parent directories
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Save to temp file first to avoid corruption/locking issues
        fd, tmp_path = tempfile.mkstemp(suffix='.docx', dir=str(output_path.parent))
        try:
            os.close(fd) # Close the file descriptor, docx will open it
            self.doc.save(tmp_path)
            # Atomic replace
            if os.path.exists(str(output_path)):
                os.remove(str(output_path))
            os.rename(tmp_path, str(output_path))
        except Exception as e:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
            logger.error(f"Failed to export DOCX: {e}")
            raise

        logger.info(f"DOCX exported successfully: {output_path}")
        return str(output_path)

    def add_table_of_contents(self):
        """Add placeholder for table of contents"""
        self.doc.add_heading('Table of Contents', level=1)
        para = self.doc.add_paragraph()
        para.add_run('[Right-click here and select "Update Field" to update TOC]')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        logger.debug("Table of contents placeholder added")


def create_simple_document(
    title: str,
    author: str,
    content: str,
    output_path: str = 'document.docx'
) -> str:
    """
    Helper function to create a simple DOCX document

    Args:
        title: Document title
        author: Document author
        content: Document content
        output_path: Output file path

    Returns:
        Path to generated DOCX file
    """
    exporter = AcademicDOCXExporter(template='ieee')

    exporter.add_title_page(title, author, "University Name")
    exporter.add_abstract("Abstract text here...")

    exporter.add_chapter(1, 'Introduction', [
        {
            'title': 'Background',
            'text': content
        }
    ])

    return exporter.export(output_path)


if __name__ == '__main__':
    # Example usage
    exporter = AcademicDOCXExporter(template='ieee')

    # Title page
    exporter.add_title_page(
        "My PhD Thesis",
        "John Doe",
        "University of Technology"
    )

    # Abstract
    exporter.add_abstract(
        "This thesis presents novel approaches to drug discovery..."
    )

    # Chapter with sections
    exporter.add_chapter(1, 'Introduction', [
        {
            'title': 'Background',
            'text': 'Drug discovery is a complex process...',
            'tables': [
                {
                    'headers': ['Method', 'Accuracy', 'Time (s)'],
                    'rows': [
                        ['Method A', '95%', '120'],
                        ['Method B', '92%', '85']
                    ],
                    'caption': 'Comparison of methods'
                }
            ]
        },
        {
            'title': 'Objectives',
            'text': 'The main objectives of this research are:'
        }
    ])

    # References
    exporter.add_references([
        "Smith, J. et al. (2023). Novel drug discovery method.",
        "Doe, A. (2022). Advanced pharmaceutical research."
    ])

    # Export
    exporter.export('thesis.docx')
